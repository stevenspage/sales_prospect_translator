# 安装依赖：
# pip install googlesearch-python newspaper3k fpdf python-docx trafilatura langdetect pandas openpyxl

# ==================== 用户配置区域 ====================
SEARCH_KEYWORD = "importador argentino de máquinas de café"  # 搜索关键词，可在此处修改

# 搜索配置
SEARCH_CONFIG = {
    'num_results': 300,  # 搜索结果数量
    'sleep_interval_min': 1,  # 最小请求间隔（秒）
    'sleep_interval_max': 2  # 最大请求间隔（秒）
}

# 正文截取配置
TEXT_TRUNCATE_CONFIG = {
    'max_chars': 500,  # 最大字符数
    'enable_truncate': True  # 是否启用正文截取
}

# QQ邮箱配置
EMAIL_CONFIG = {
    'smtp_server': 'smtp.qq.com',
    'smtp_port': 587,
    'sender_email': '',  # 请替换为您的QQ邮箱
    'sender_password': '',  # 请替换为您的QQ邮箱授权码
    'recipient_email': '',  # 请替换为收件人邮箱
    'enable_email': False  # 是否启用邮件发送功能
}

# 智谱翻译API配置
TRANSLATION_CONFIG = {
    'api_url': 'https://open.bigmodel.cn/api/v1/agents',
    'agent_id': 'general_translation',
    'api_key': '',  # 请替换为您的API密钥
    'source_lang': 'auto',    # 自动检测源语言
    'target_lang': 'zh-CN',   # 目标语言：简体中文
    'strategy': 'two_step',   # 翻译策略：两步翻译
    'enable_translation': True,  # 是否启用翻译功能
    'translate_content': True,   # 是否翻译正文内容
    'translate_summary': True,   # 是否翻译摘要
    'max_workers': 10,          # 并发翻译线程数（建议3-5个）
    'request_delay': 0.5       # 请求间隔（秒），避免API限流
}
# =====================================================

from googlesearch import search
from newspaper import Article
from requests.exceptions import ReadTimeout
from fpdf import FPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests
from bs4 import BeautifulSoup
import html
from urllib.parse import urlparse
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import random
import concurrent.futures
import threading
import time
import pandas as pd

# 多库组合导入
try:
    import trafilatura
    TRAFILATURA_AVAILABLE = True
except ImportError:
    TRAFILATURA_AVAILABLE = False
    print("警告: trafilatura 未安装，将跳过此库")

# 语言检测导入
try:
    from langdetect import detect, detect_langs
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False
    print("警告: langdetect 未安装，将使用默认英语搜索")

# 语言映射配置
LANGUAGE_MAPPING = {
    'zh-cn': 'zh-CN',  # 中文
    'en': 'en',        # 英语
    'ja': 'ja',        # 日语
    'es': 'es',        # 西班牙语
    'de': 'de',        # 德语
    'fr': 'fr',        # 法语
    'pt': 'pt',        # 葡萄牙语
    'ko': 'ko',        # 韩语
}

def detect_search_language(keyword):
    """
    检测搜索关键词的语言
    """
    if not LANGDETECT_AVAILABLE:
        return 'en'  # 默认英语
    
    try:
        # 检测语言
        detected_lang = detect(keyword)
        
        # 映射到Google搜索语言代码
        google_lang = LANGUAGE_MAPPING.get(detected_lang, 'en')
        
        return google_lang
        
    except Exception as e:
        return 'en'

# 创建线程锁，用于控制并发请求
translation_lock = threading.Lock()

def translate_text(text, text_type="content"):
    """
    使用智谱翻译API翻译文本（支持并发）
    """
    if not TRANSLATION_CONFIG['enable_translation'] or not text.strip():
        return text
    
    # 检查是否需要翻译
    if text_type == "content" and not TRANSLATION_CONFIG['translate_content']:
        return text
    if text_type == "summary" and not TRANSLATION_CONFIG['translate_summary']:
        return text
    
    try:
        # 添加请求延迟，避免API限流
        time.sleep(TRANSLATION_CONFIG['request_delay'])
        
        # 直接使用传入的文本进行翻译（文本已在调用前截取）
        text_to_translate = text
        
        payload = {
            "agent_id": TRANSLATION_CONFIG['agent_id'],
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": text_to_translate
                        }
                    ]
                }
            ],
            "custom_variables": {
                "source_lang": TRANSLATION_CONFIG['source_lang'],
                "target_lang": TRANSLATION_CONFIG['target_lang'],
                "strategy": TRANSLATION_CONFIG['strategy']
            }
        }
        
        headers = {
            "Authorization": f"Bearer {TRANSLATION_CONFIG['api_key']}",
            "Content-Type": "application/json"
        }
        
        with translation_lock:
            print(f"  正在翻译{text_type}...")
        
        response = requests.post(TRANSLATION_CONFIG['api_url'], json=payload, headers=headers, timeout=30)
        response.raise_for_status()
        
        result = response.json()
        
        if 'choices' in result and len(result['choices']) > 0:
            content = result['choices'][0]['messages'][0]['content']
            # 根据API文档，content是对象，包含type和text字段
            if isinstance(content, dict) and 'text' in content:
                translated_text = content['text']
                with translation_lock:
                    print(f"  {text_type}翻译完成")
                return translated_text
            elif isinstance(content, str):
                # 兼容直接返回字符串的情况
                translated_text = content
                with translation_lock:
                    print(f"  {text_type}翻译完成")
                return translated_text
            else:
                with translation_lock:
                    print(f"  {text_type}翻译失败：content格式异常")
                return text
        else:
            with translation_lock:
                print(f"  {text_type}翻译失败：API返回格式异常")
            return text
            
    except requests.exceptions.Timeout:
        with translation_lock:
            print(f"  {text_type}翻译超时，返回原文")
        return text
    except requests.exceptions.RequestException as e:
        with translation_lock:
            print(f"  {text_type}翻译请求失败：{e}")
        return text
    except Exception as e:
        with translation_lock:
            print(f"  {text_type}翻译异常：{e}")
        return text

def translate_single_result(result_data):
    """
    翻译单个搜索结果（用于并发执行）
    """
    i, url, title, text, description = result_data
    
    print(f"\n--- 翻译第 {i} 个结果 ---")
    print(f"URL: {url}")
    
    # 翻译标题
    translated_title = translate_text(title, "title")
    
    # 翻译摘要
    translated_description = translate_text(description, "summary")
    
    # 先截取正文到500字符，然后翻译截取后的文本
    if TEXT_TRUNCATE_CONFIG['enable_truncate']:
        truncated_text = truncate_text_smart(text, TEXT_TRUNCATE_CONFIG['max_chars'])
        print(f"  正文已截取到 {len(truncated_text)} 字符")
        # 翻译截取后的文本
        translated_text = translate_text(truncated_text, "content")
    else:
        # 如果没有启用截取，仍然截取到500字符进行翻译
        truncated_text = truncate_text_smart(text, 500)
        print(f"  正文已截取到 {len(truncated_text)} 字符进行翻译")
        translated_text = translate_text(truncated_text, "content")
    
    print(f"第 {i} 个结果翻译完成")
    return (url, title, text, description, translated_title, translated_text, translated_description)

def batch_translate_results(results):
    """
    批量翻译搜索结果（并发翻译）
    """
    if not TRANSLATION_CONFIG['enable_translation']:
        print("翻译功能已禁用，跳过翻译")
        return results
    
    print(f"\n开始并发翻译 {len(results)} 个搜索结果...")
    print(f"并发线程数: {TRANSLATION_CONFIG['max_workers']}")
    print(f"请求间隔: {TRANSLATION_CONFIG['request_delay']}秒")
    
    # 准备并发翻译的数据
    translation_tasks = [(i+1, url, title, text, description) for i, (url, title, text, description) in enumerate(results)]
    
    translated_results = []
    
    # 使用线程池进行并发翻译
    with concurrent.futures.ThreadPoolExecutor(max_workers=TRANSLATION_CONFIG['max_workers']) as executor:
        # 提交所有翻译任务
        future_to_result = {executor.submit(translate_single_result, task): task for task in translation_tasks}
        
        # 收集翻译结果
        for future in concurrent.futures.as_completed(future_to_result):
            try:
                result = future.result()
                translated_results.append(result)
            except Exception as e:
                print(f"翻译任务失败: {e}")
                # 如果翻译失败，添加原始数据
                original_task = future_to_result[future]
                translated_results.append((original_task[1], original_task[2], original_task[3], original_task[4], 
                                         original_task[2], original_task[3], original_task[4]))
    
    # 按原始顺序排序结果
    translated_results.sort(key=lambda x: next(i for i, task in enumerate(translation_tasks) if task[1] == x[0]))
    
    print(f"\n并发翻译完成！共翻译了 {len(translated_results)} 个结果")
    return translated_results

def get_google_results(query, num_results=None, sleep_interval=None):
    """
    获取 Google 搜索结果前 num_results 个结果（包含URL、标题和摘要）
    支持自动语言检测和随机请求间隔控制
    """
    # 使用配置中的默认值
    if num_results is None:
        num_results = SEARCH_CONFIG['num_results']
    
    # 自动检测搜索语言
    detected_lang = detect_search_language(query)
    
    # 如果没有指定sleep_interval，使用配置中的随机值
    if sleep_interval is None:
        sleep_interval = random.uniform(SEARCH_CONFIG['sleep_interval_min'], SEARCH_CONFIG['sleep_interval_max'])
        print(f"使用随机请求间隔: {sleep_interval:.1f}秒")
    
    results = []
    try:
        for result in search(query, num_results=num_results, lang=detected_lang, sleep_interval=sleep_interval, advanced=True):
            # 获取摘要信息，如果没有则使用空字符串
            description = getattr(result, 'description', '') or ''
            results.append((result.url, result.title, description))
    except Exception as e:
        print(f"获取搜索结果失败: {e}")
    return results

def extract_with_newspaper(url):
    """
    使用newspaper3k提取正文
    """
    try:
        article = Article(url)
        article.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        article.download()
        article.parse()
        return article.text
    except Exception as e:
        print(f"newspaper3k提取失败: {url}, 错误: {e}")
        return ""

def extract_with_trafilatura(url):
    """
    使用trafilatura提取正文
    """
    if not TRAFILATURA_AVAILABLE:
        return ""
    
    try:
        downloaded = trafilatura.fetch_url(url)
        if downloaded:
            content = trafilatura.extract(downloaded)
            return content if content else ""
        return ""
    except Exception as e:
        print(f"trafilatura提取失败: {url}, 错误: {e}")
        return ""



def extract_article_text_robust(url):
    """
    多库组合提取正文 - 主要方法
    """
    # 定义提取方法列表，按优先级排序
    extractors = [
        ("trafilatura", extract_with_trafilatura),
        ("newspaper3k", extract_with_newspaper)
    ]
    
    best_content = ""
    best_method = ""
    best_length = 0
    
    print(f"开始多库组合提取: {url}")
    
    for method_name, extractor in extractors:
        try:
            print(f"  尝试 {method_name}...")
            content = extractor(url)
            
            if content and len(content) > 100:  # 至少100个字符
                content_length = len(content)
                print(f"  {method_name} 成功，提取了 {content_length} 个字符")
                
                # 选择最长的内容作为最佳结果
                if content_length > best_length:
                    best_content = content
                    best_method = method_name
                    best_length = content_length
                    
                    # 如果内容足够长，可以提前结束
                    if content_length > 2000:
                        print(f"  内容足够长，使用 {method_name} 的结果")
                        break
            else:
                print(f"  {method_name} 提取内容太短或为空")
                
        except Exception as e:
            print(f"  {method_name} 提取异常: {e}")
            continue
    
    if best_content:
        print(f"最终选择: {best_method}，内容长度: {len(best_content)}")
        return best_content
    else:
        print("所有提取方法都失败了")
        return ""

def extract_article_text(url):
    """
    提取网页正文 - 保持向后兼容
    """
    return extract_article_text_robust(url)

def truncate_text_smart(text, max_chars=500):
    """
    智能截取文本，保持句子完整性
    """
    if not text or len(text) <= max_chars:
        return text
    
    # 截取到指定长度
    truncated = text[:max_chars]
    
    # 查找最后一个完整的句子结束符
    sentence_endings = ['.', '!', '?', '。', '！', '？']
    last_sentence_end = -1
    
    for ending in sentence_endings:
        pos = truncated.rfind(ending)
        if pos > last_sentence_end:
            last_sentence_end = pos
    
    # 如果找到句子结束符，在它后面截断
    if last_sentence_end > max_chars * 0.7:  # 确保不会截掉太多内容
        return truncated[:last_sentence_end + 1]
    else:
        # 如果没找到合适的句子结束符，在单词边界截断
        last_space = truncated.rfind(' ')
        if last_space > max_chars * 0.8:
            return truncated[:last_space] + "..."
        else:
            return truncated + "..."

def extract_meta_from_html(url):
    """
    从HTML中提取标题和description
    """
    try:
        # 设置请求头模拟真实浏览器
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # 提取标题
        title = ""
        title_tag = soup.find('title')
        if title_tag:
            title = title_tag.get_text()
            title = html.unescape(title)
            import re
            title = re.sub(r'\s+', ' ', title).strip()
        
        # 提取description
        description = ""
        # 尝试多种description标签
        desc_selectors = [
            'meta[name="description"]',
            'meta[property="og:description"]',
            'meta[name="twitter:description"]'
        ]
        
        for selector in desc_selectors:
            desc_tag = soup.select_one(selector)
            if desc_tag and desc_tag.get('content'):
                description = desc_tag.get('content')
                description = html.unescape(description)
                description = re.sub(r'\s+', ' ', description).strip()
                break
        
        return title, description
    except Exception as e:
        print(f"提取元数据失败: {url}, 错误: {e}")
        return "", ""

def extract_article_title_and_text(url, google_title="", google_description=""):
    """
    提取网页标题、description和正文（使用多库组合），不进行翻译
    """
    try:
        # 从HTML获取标题和description
        html_title, html_description = extract_meta_from_html(url)
        
        # 使用多库组合获取正文
        text = extract_article_text_robust(url)
        
        # 确定最终使用的标题
        title = html_title if html_title else google_title
        if not title and google_title:
            title = google_title
            print(f"使用Google搜索结果标题作为备选: {title}")
        
        # 确定最终使用的description（优先级：HTML > Google）
        final_description = html_description if html_description else google_description
        
        return title, text, final_description
    except Exception as e:
        print(f"提取失败: {url}, 错误: {e}")
        return "", "", ""

def save_results_to_pdf(results, filename="results.pdf"):
    """
    将搜索结果保存为 PDF 文件，支持中文
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # 添加支持中文的字体
    pdf.add_font('ArialUnicode', '', 'msyh.ttf', uni=True)
    pdf.set_font('ArialUnicode', size=12)

    for i, (url, title, text) in enumerate(results, 1):
        pdf.set_font('ArialUnicode', style='B', size=12)
        pdf.cell(0, 10, f"第 {i} 个搜索结果", ln=True)
        pdf.set_font('ArialUnicode', size=12)
        pdf.cell(0, 10, f"标题: {title}", ln=True)
        pdf.cell(0, 10, f"URL: {url}", ln=True)
        pdf.multi_cell(0, 10, f"正文内容: {text}")
        pdf.ln(10)

    pdf.output(filename)

def add_hyperlink(paragraph, text, url):
    """
    在段落中添加超链接，并设置蓝色字体和下划线样式
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    # 设置蓝色字体和下划线样式
    color = OxmlElement("w:color")
    color.set("val", "0000FF")  # 蓝色
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set("val", "single")  # 下划线
    r_pr.append(underline)

    new_run.append(r_pr)

    new_text = OxmlElement("w:t")
    new_text.text = text
    new_run.append(new_text)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def extract_domain(url):
    """
    从URL中提取域名
    """
    try:
        parsed_url = urlparse(url)
        domain = parsed_url.netloc
        # 移除www前缀
        if domain.startswith('www.'):
            domain = domain[4:]
        return domain
    except Exception:
        return url

def save_results_to_docx(results, filename="results.docx"):
    """
    将搜索结果保存为 DOCX 文件，标题添加超链接，支持翻译内容
    """
    doc = Document()
    doc.add_heading(f"搜索结果 - {SEARCH_KEYWORD} ({__import__('datetime').datetime.now().strftime('%Y-%m-%d')})", level=1)
    doc.add_paragraph(f"搜索语种: {detect_search_language(SEARCH_KEYWORD).upper()}")

    for i, result in enumerate(results, 1):
        # 解包结果，支持翻译内容
        if len(result) == 7:  # 包含翻译内容
            url, title, text, description, translated_title, translated_text, translated_description = result
        else:  # 不包含翻译内容（向后兼容）
            url, title, text, description = result
            translated_title, translated_text, translated_description = title, text, description
        
        doc.add_heading(f"第 {i} 个搜索结果", level=2)
        
        # 添加标题（原始标题带超链接，翻译标题紧跟其后）
        title_paragraph = doc.add_paragraph()
        add_hyperlink(title_paragraph, title, url)
        
        # 添加翻译后标题（如果启用翻译），使用相同样式
        if TRANSLATION_CONFIG['enable_translation'] and translated_title != title and translated_title:
            # 在同一个段落中添加翻译标题
            title_paragraph.add_run(f" {translated_title}")
        
        # 添加URL信息
        doc.add_paragraph(f"网页URL: {url}")
        
        # 添加公司网址信息
        domain = extract_domain(url)
        doc.add_paragraph(f"公司网址: {domain}")
        
        # 添加网站摘要信息
        if description:
            doc.add_paragraph(f"摘要: {description}")
        
        # 添加翻译后摘要（如果启用翻译）
        if TRANSLATION_CONFIG['enable_translation'] and translated_description != description and translated_description:
            doc.add_paragraph(translated_description)
        
        # 添加正文内容（智能截取）
        doc.add_paragraph("正文内容:")
        if TEXT_TRUNCATE_CONFIG['enable_truncate']:
            truncated_text = truncate_text_smart(text, TEXT_TRUNCATE_CONFIG['max_chars'])
            doc.add_paragraph(truncated_text)
        else:
            doc.add_paragraph(text)
        
        # 添加翻译后正文内容（如果启用翻译）
        if TRANSLATION_CONFIG['enable_translation'] and translated_text != text and translated_text:
            if TEXT_TRUNCATE_CONFIG['enable_truncate']:
                truncated_translated_text = truncate_text_smart(translated_text, TEXT_TRUNCATE_CONFIG['max_chars'])
                doc.add_paragraph(truncated_translated_text)
            else:
                doc.add_paragraph(translated_text)

    doc.save(filename)

def save_results_to_html(results, filename="results.html"):
    """
    将搜索结果保存为简洁的HTML文件
    """
    html_content = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>搜索结果 - {SEARCH_KEYWORD}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            background: #fff;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        
        h1 {{
            font-size: 2em;
            font-weight: 600;
            margin-bottom: 20px;
            color: #000;
            border-bottom: 2px solid #e1e4e8;
            padding-bottom: 10px;
        }}
        
        .search-info {{
            background: #f6f8fa;
            border: 1px solid #d0d7de;
            border-radius: 6px;
            padding: 16px;
            margin-bottom: 30px;
            font-size: 14px;
            color: #656d76;
        }}
        
        .result-item {{
            margin-bottom: 40px;
            padding-bottom: 30px;
            border-bottom: 1px solid #e1e4e8;
        }}
        
        .result-item:last-child {{
            border-bottom: none;
        }}
        
        .result-number {{
            display: inline-block;
            background: #0969da;
            color: white;
            width: 24px;
            height: 24px;
            border-radius: 12px;
            text-align: center;
            line-height: 24px;
            font-size: 12px;
            font-weight: 600;
            margin-right: 12px;
            vertical-align: top;
        }}
        
        .result-title {{
            font-size: 1.25em;
            font-weight: 600;
            margin-bottom: 8px;
            line-height: 1.3;
        }}
        
        .result-title a {{
            color: #0969da;
            text-decoration: none;
        }}
        
        .result-title a:hover {{
            text-decoration: underline;
        }}
        
        .result-title-link {{
            color: #0969da;
            text-decoration: none;
        }}
        
        .result-title-link:hover {{
            text-decoration: underline;
        }}
        
        .result-url {{
            color: #656d76;
            font-size: 14px;
            margin-bottom: 8px;
            word-break: break-all;
        }}
        
        .result-url a {{
            color: #0969da;
            text-decoration: none;
        }}
        
        .result-url a:hover {{
            text-decoration: underline;
        }}
        
        .domain-badge {{
            display: inline-block;
            background: #dbeafe;
            color: #1e40af;
            padding: 2px 6px;
            border-radius: 3px;
            font-size: 12px;
            font-weight: 500;
            margin-left: 8px;
        }}
        
        .result-content {{
            margin-top: 12px;
            font-size: 14px;
            line-height: 1.5;
            color: #656d76;
        }}
        
        .result-content h4 {{
            color: #24292f;
            font-size: 14px;
            font-weight: 600;
            margin: 16px 0 8px 0;
        }}
        
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #e1e4e8;
            font-size: 12px;
            color: #656d76;
            text-align: center;
        }}
        
        @media (max-width: 768px) {{
            body {{
                padding: 16px;
            }}
            
            h1 {{
                font-size: 1.5em;
            }}
            
            .result-title {{
                font-size: 1.1em;
            }}
        }}
    </style>
</head>
<body>
    <h1>搜索结果-{SEARCH_KEYWORD} ({__import__('datetime').datetime.now().strftime('%Y-%m-%d')})</h1>
    
    <div class="search-info">
        <strong>搜索语种:</strong> {detect_search_language(SEARCH_KEYWORD).upper()} | 
        <strong>结果数量:</strong> {len(results)} 个
    </div>
"""

    for i, result in enumerate(results, 1):
        # 解包结果，支持翻译内容
        if len(result) == 7:  # 包含翻译内容
            url, title, text, description, translated_title, translated_text, translated_description = result
        else:  # 不包含翻译内容（向后兼容）
            url, title, text, description = result
            translated_title, translated_text, translated_description = title, text, description
        
        domain = extract_domain(url)
        html_content += f"""
    <div class="result-item">
        <div class="result-title">
            <span class="result-number">{i}</span>
            <a href="{url}" target="_blank">{title}</a>
            {f' <a href="{url}" target="_blank" class="result-title-link">{translated_title}</a>' if TRANSLATION_CONFIG['enable_translation'] and translated_title != title and translated_title else ''}
        </div>
        <div class="result-url">
            <a href="{url}" target="_blank">{url}</a>
            <span class="domain-badge">{domain}</span>
        </div>
        {f'<div class="result-content"><p><strong>摘要:</strong> {description}</p></div>' if description else ''}
        {f'<div class="result-content"><p>{translated_description}</p></div>' if TRANSLATION_CONFIG['enable_translation'] and translated_description != description and translated_description else ''}
        <div class="result-content">
            <h4>正文内容:</h4>
            <p>{truncate_text_smart(text, TEXT_TRUNCATE_CONFIG['max_chars']) if TEXT_TRUNCATE_CONFIG['enable_truncate'] else text}</p>
        </div>
        {f'<div class="result-content"><p>{truncate_text_smart(translated_text, TEXT_TRUNCATE_CONFIG["max_chars"]) if TEXT_TRUNCATE_CONFIG["enable_truncate"] else translated_text}</p></div>' if TRANSLATION_CONFIG['enable_translation'] and translated_text != text and translated_text else ''}
    </div>
"""

    html_content += f"""
    <div class="footer">
        <p>搜索结果由 Google 搜索生成 | 生成时间: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
</body>
</html>
"""

    with open(filename, 'w', encoding='utf-8') as f:
        f.write(html_content)

def save_results_to_excel(results, filename="results.xlsx"):
    """
    将搜索结果保存为 Excel 文件
    """
    # 准备数据列表
    data = []
    
    for i, result in enumerate(results, 1):
        # 解包结果，支持翻译内容
        if len(result) == 7:  # 包含翻译内容
            url, title, text, description, translated_title, translated_text, translated_description = result
        else:  # 不包含翻译内容（向后兼容）
            url, title, text, description = result
            translated_title, translated_text, translated_description = title, text, description
        
        # 提取域名
        domain = extract_domain(url)
        
        # 截取正文内容
        if TEXT_TRUNCATE_CONFIG['enable_truncate']:
            truncated_text = truncate_text_smart(text, TEXT_TRUNCATE_CONFIG['max_chars'])
            truncated_translated_text = truncate_text_smart(translated_text, TEXT_TRUNCATE_CONFIG['max_chars']) if TRANSLATION_CONFIG['enable_translation'] else ""
        else:
            truncated_text = text
            truncated_translated_text = translated_text if TRANSLATION_CONFIG['enable_translation'] else ""
        
        # 添加数据行
        row_data = {
            '序号': i,
            '原始标题': title,
            '翻译标题': translated_title if TRANSLATION_CONFIG['enable_translation'] and translated_title != title else "",
            '网页URL': url,
            '公司网址': domain,
            '原始摘要': description,
            '翻译摘要': translated_description if TRANSLATION_CONFIG['enable_translation'] and translated_description != description else "",
            '原始正文': truncated_text,
            '翻译正文': truncated_translated_text if TRANSLATION_CONFIG['enable_translation'] and translated_text != text else ""
        }
        data.append(row_data)
    
    # 创建DataFrame
    df = pd.DataFrame(data)
    
    # 保存为Excel文件
    with pd.ExcelWriter(filename, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='搜索结果', index=False)
        
        # 获取工作表对象以调整列宽
        worksheet = writer.sheets['搜索结果']
        
        # 自动调整列宽
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)  # 限制最大宽度为50
            worksheet.column_dimensions[column_letter].width = adjusted_width

def send_email_with_attachment(html_filename, docx_filename, xlsx_filename, subject="搜索结果报告", num_results=0):
    """
    通过QQ邮箱发送HTML内容作为邮件正文，并附加Word文档和Excel文件
    """
    if not EMAIL_CONFIG['enable_email']:
        print("邮件发送功能已禁用，请在配置中启用。")
        return False
    
    try:
        # 读取HTML文件内容
        with open(html_filename, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # 创建邮件对象
        msg = MIMEMultipart('mixed')
        msg['From'] = EMAIL_CONFIG['sender_email']
        msg['To'] = EMAIL_CONFIG['recipient_email']
        msg['Subject'] = f"{subject} - {SEARCH_KEYWORD} ({__import__('datetime').datetime.now().strftime('%Y-%m-%d')})"
        
        # 纯文本版本
        text_body = f"""
        您好！
        
        这是关于"{SEARCH_KEYWORD}"的搜索结果报告。
        
        报告包含：
        - {num_results}个相关网站的详细信息
        - 网站标题、URL、域名和摘要
        - 完整的网页正文内容
        - 中文翻译版本
        
        生成时间：{__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        
        此邮件由自动搜索脚本生成。
        
        附件说明：
        - Word文档：包含完整的搜索结果和翻译，适合阅读和打印
        - Excel文件：包含结构化的数据表格，适合数据分析和处理
        
        注意：此邮件包含HTML格式的搜索结果和多个附件，请确保您的邮件客户端支持HTML显示和附件下载。
        """
        
        # 创建纯文本和HTML部分
        text_part = MIMEText(text_body, 'plain', 'utf-8')
        html_part = MIMEText(html_content, 'html', 'utf-8')
        
        # 创建多部分容器
        msg_alternative = MIMEMultipart('alternative')
        msg_alternative.attach(text_part)
        msg_alternative.attach(html_part)
        msg.attach(msg_alternative)
        
        # 添加Word文档附件
        try:
            with open(docx_filename, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename= {docx_filename}'
                )
                msg.attach(part)
            print(f"Word文档附件已添加: {docx_filename}")
        except FileNotFoundError:
            print(f"警告: 找不到Word文档文件 {docx_filename}")
        except Exception as e:
            print(f"添加Word文档附件失败: {e}")
        
        # 添加Excel文件附件
        try:
            with open(xlsx_filename, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename= {xlsx_filename}'
                )
                msg.attach(part)
            print(f"Excel文件附件已添加: {xlsx_filename}")
        except FileNotFoundError:
            print(f"警告: 找不到Excel文件 {xlsx_filename}")
        except Exception as e:
            print(f"添加Excel文件附件失败: {e}")
        
        # 连接SMTP服务器并发送邮件
        server = smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port'])
        server.starttls()  # 启用TLS加密
        server.login(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['sender_password'])
        
        text = msg.as_string()
        server.sendmail(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['recipient_email'], text)
        server.quit()
        
        print(f"邮件发送成功！收件人：{EMAIL_CONFIG['recipient_email']}")
        print("HTML内容已作为邮件正文发送，Word文档和Excel文件已作为附件发送。")
        return True
        
    except Exception as e:
        print(f"邮件发送失败：{e}")
        return False

def main():
    query = SEARCH_KEYWORD
    print(f"搜索关键词: {query}")
    search_results = get_google_results(query)

    if not search_results:
        print("未能获取任何搜索结果。")
        return

    # 第一阶段：提取所有搜索结果的内容（不翻译）
    print(f"\n{'='*50}")
    print("第一阶段：提取搜索结果内容")
    print(f"{'='*50}")
    
    results = []
    for i, (url, google_title, google_description) in enumerate(search_results, 1):
        print(f"\n=== 第 {i} 个搜索结果 ===")
        print("URL:", url)
        print("Google标题:", google_title)
        print("Google摘要:", google_description)
        
        # 只提取内容，不翻译
        title, text, description = extract_article_title_and_text(url, google_title, google_description)
        
        print("最终标题:", title)
        print("最终摘要:", description)
        print("正文内容（智能截取500字）：")
        if TEXT_TRUNCATE_CONFIG['enable_truncate']:
            truncated_text = truncate_text_smart(text, TEXT_TRUNCATE_CONFIG['max_chars'])
            print(truncated_text)
        else:
            print(text)
        
        # 保存原始信息到结果列表
        results.append((url, title, text, description))

    # 第二阶段：统一翻译所有结果
    print(f"\n{'='*50}")
    print("第二阶段：批量翻译搜索结果")
    print(f"{'='*50}")
    
    # 批量翻译所有结果
    translated_results = batch_translate_results(results)

    # 第三阶段：保存结果
    print(f"\n{'='*50}")
    print("第三阶段：保存结果")
    print(f"{'='*50}")
    
    # 根据搜索关键词生成文件名
    safe_filename = "".join(c for c in SEARCH_KEYWORD if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_filename = safe_filename.replace(' ', '_')
    docx_filename = f"{safe_filename}.docx"
    html_filename = f"{safe_filename}.html"
    xlsx_filename = f"{safe_filename}.xlsx"
    
    # 保存为Word文档
    save_results_to_docx(translated_results, docx_filename)
    print(f"Word文档已保存到 {docx_filename} 文件中。")
    
    # 保存为HTML文件
    save_results_to_html(translated_results, html_filename)
    print(f"HTML文件已保存到 {html_filename} 文件中。")
    
    # 保存为Excel文件
    save_results_to_excel(translated_results, xlsx_filename)
    print(f"Excel文件已保存到 {xlsx_filename} 文件中。")
    
    # 发送邮件（如果启用）
    if EMAIL_CONFIG['enable_email']:
        send_email_with_attachment(html_filename, docx_filename, xlsx_filename, num_results=len(translated_results))

if __name__ == "__main__":
    main()