# ==================== 1. Google搜索配置组 ====================
# Google官方API 密钥配置（必须配置）
# 请替换为您的Google官方API密钥，例如GOOGLE_API_KEY = 'AIza**************'
GOOGLE_API_KEY = ''  

# 请替换为您的Google可编程搜索引擎ID，例如GOOGLE_SEARCH_ENGINE_ID = '14d0*****'
GOOGLE_SEARCH_ENGINE_ID = '' 

# Serp API 密钥配置（使用Serp API时必填）
# 请替换为您的Serp API密钥
SERP_API_KEY = ''

# 搜索关键词配置（替换为您的搜索关键词）
SEARCH_KEYWORD = 'distribuidor cámaras de acción Argentina'

# 排除关键词配置（用空格分隔多个关键词，如 '垃圾 广告 推广'）
EXCLUDE_TERMS = ''

# 搜索方法选择-填写范例：
# SEARCH_METHOD = 'google_api'                # 使用Google官方API
# SEARCH_METHOD = 'googlesearch'              # 使用googlesearch-python库
# SEARCH_METHOD = 'serp_api'                  # 使用Serp API（第三方搜索API）
SEARCH_METHOD = 'googlesearch'

# ==================== 2. 搜索参数配置组 ====================
# 搜索配置
SEARCH_CONFIG = {
    # 基本搜索设置
    'num_results': 20,  # 搜索结果数量
    # Google Custom Search  设置
    'language': 'Google默认',  # 搜索结果语言：'auto' （根据搜索关键词自动检测，支持主要语言，对其他语言识别可能不精确）
                         # 'Google默认' （Google默认设置，不限制语言）
                         # 常见语言代码：
                         # 'lang_en'      - 英语
                         # 'lang_zh-CN'   - 简体中文
                         # 'lang_zh-TW'   - 繁体中文
                         # 'lang_es'      - 西班牙语
                         # 'lang_ja'      - 日语
                         # 'lang_ko'      - 韩语
                         # 'lang_ru'      - 俄语
                         # 'lang_vi'      - 越南语
                         # 'lang_ar'      - 阿拉伯语
                         # 'lang_de'      - 德语
                         # 'lang_fr'      - 法语
                         # 'lang_pt'      - 葡萄牙语
    'country': 'Google默认',  # 搜索国家：填入下列国家代码，如 US 或 ES
                         # 注意：Google官方API和Serp API使用cr参数（countryUS格式），googlesearch-python使用region参数（小写格式）
                         # 常见国家代码（ISO 3166-1 国家代码,https://www.alibabacloud.com/help/zh/chatapp/country-or-region-codes）：
                         # 'CN'        - 中国
                         # 'US'        - 美国
                         # 'GB'        - 英国
                         # 'JP'        - 日本
                         # 'KR'        - 韩国
                         # 'IN'        - 印度
                         # 'FR'        - 法国
                         # 'DE'        - 德国
                         # 'RU'        - 俄罗斯
                         # 'ES'        - 西班牙
                         # 'BR'        - 巴西
                         # 'AR'        - 阿根廷
                         # 'MX'        - 墨西哥
                         # 'VN'        - 越南
}

# 统一语言配置常量
# 1) 语言代码数据源（生成"搜索语言"工作表与下拉）
LANGUAGE_OPTIONS = [
    'Google默认', 'auto','lang_en', 'lang_zh-CN', 'lang_zh-TW', 'lang_es', 'lang_ja', 'lang_ko', 'lang_ru', 'lang_vi', 'lang_ar', 'lang_de', 'lang_fr', 'lang_pt'
]

# 2) 语言中文显示映射（用于 UI 显示）
LANGUAGE_DISPLAY_MAP = {
    'en': '英语',
    'zh-cn': '简体中文',
    'zh-tw': '繁体中文',
    'es': '西班牙语',
    'ja': '日语',
    'ko': '韩语',
    'ru': '俄语',
    'vi': '越南语',
    'ar': '阿拉伯语',
    'de': '德语',
    'fr': '法语',
    'pt': '葡萄牙语',
}

# 3) 自动检测到 Google lr 代码的特殊映射
SPECIAL_LANG_PREFIX_MAP = {
    'zh-cn': 'lang_zh-CN',
    'zh-tw': 'lang_zh-TW',
    'he': 'lang_iw'
}


# ==================== 通用标准化与解析工具 ====================
def normalize_cell(value: object) -> str:
    """
    将单元格值标准化为字符串：
    - None/NaN/空白/'nan'/'none'/'null' -> ''
    - 其余转为去首尾空格的字符串
    """
    try:
        import pandas as _pd  # 局部导入以避免顶层依赖顺序问题
        if value is None or (isinstance(value, float) and _pd.isna(value)):
            return ''
    except Exception:
        if value is None:
            return ''
    s = str(value).strip()
    return '' if s.lower() in ['nan','Google默认', 'none', 'null'] else s

def parse_int_or_none(value: object):
    # 尝试将值解析为 int；无法解析或为空时返回 None。
    s = normalize_cell(value)
    if s == '':
        return None
    try:
        return int(float(s)) if ('.' in s) else int(s)
    except Exception:
        return None

# 其他搜索配置
OTHER_CONFIG = {
    'sleep_interval': 3,  # 搜索请求间隔时间（秒），Google API和SERP API使用固定值
}

# 多关键词模式：硬编码总开关（True=启用批量模式；False=按单一关键词模式）
ENABLE_MULTI_KEYWORDS_MODE_DEFAULT = False
# 运行时总开关：默认等于硬编码值，若Excel存在有效配置则以Excel为准
ENABLE_MULTI_KEYWORDS_MODE = ENABLE_MULTI_KEYWORDS_MODE_DEFAULT

# ==================== 4. 正文提取配置组 ====================
# 正文提取配置（包含内容提取和文本截取相关设置）
CONTENT_EXTRACTION_CONFIG = {
    'extraction_workers': 10,  # 内容提取并发线程数（建议3-10个）
    'enable_truncate': True,  # 是否启用正文截取，True为启用，False为不启用
    'max_chars': 500,  # 最大字符数（保持默认值）
    'network_timeout': 15   # 网络请求超时时间（秒）
}

# ==================== 5. 翻译功能配置组 ====================
# 智谱翻译API配置
TRANSLATION_CONFIG = {
    'enable_translation': False,  # 是否启用翻译功能，True为启用，False为不启用
    'api_key': '',  # 请替换为您的API密钥
    'api_url': 'https://open.bigmodel.cn/api/v1/agents',
    'agent_id': 'general_translation',
    'source_lang': 'auto',    # 自动检测源语言
    'target_lang': 'zh-CN',   # 目标语言：简体中文
    'strategy': 'two_step',   # 翻译策略：两步翻译
    'translate_content': True,   # 是否翻译正文内容
    'translate_summary': True,   # 是否翻译摘要
    'max_workers': 10,          # 并发翻译线程数（建议3-10个）
    'request_delay': 0.5       # 请求时间间隔（秒），避免API限流
}

# ==================== 6. 邮件发送配置组 ====================
# QQ邮箱配置
EMAIL_CONFIG = {
    'enable_email': False,  # 是否启用邮件发送功能,True为启用，False为不启用   
    'smtp_server': 'smtp.qq.com',
    'smtp_port': 587,
    'sender_email': '',  # 请替换为您的QQ邮箱
    'sender_auth_code': '',  # 请替换为您的QQ邮箱授权码
    'recipient_email': '' # 请替换为收件人邮箱
}


# =====================================================

# googlesearch库已移植到本脚本中
from newspaper import Article
from fpdf import FPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests
from bs4 import BeautifulSoup
import html
from urllib.parse import urlparse, unquote
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import random
import concurrent.futures
import threading
import time
from time import sleep
import pandas as pd
import os
import signal
import sys
import atexit

# Serp API 导入
try:
    from serpapi import GoogleSearch
    SERPAPI_AVAILABLE = True
except ImportError:
    SERPAPI_AVAILABLE = False
    print("警告: serpapi 未安装，将跳过Serp API搜索方法")

# 多库组合导入
try:
    import trafilatura
    TRAFILATURA_AVAILABLE = True
except ImportError:
    TRAFILATURA_AVAILABLE = False
    print("警告: trafilatura 未安装，将跳过此库")

# 语言检测导入
try:
    from langdetect import detect, detect_langs
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False
    print("警告: langdetect 未安装，将使用默认英语搜索")

# （已统一到 LANGUAGE_OPTIONS / LANGUAGE_DISPLAY_MAP / SPECIAL_LANG_PREFIX_MAP，不再需要额外映射表）

# ==================== googlesearch库代码 ====================
# 以下代码来自googlesearch库，已移植到本脚本中

def get_useragent():
    # Lynx was deprecated 2025.9, so we using something else.
    return f"AdsBot-Google (+http://www.google.com/adsbot.html)"


def _req(q, results, start, proxies, timeout, safe, ssl_verify, gl, hl=None, cr=None, lr=None):
    # 构建参数字典 - 所有参数都始终存在，None表示Google默认设置
    params = {
        "q": q,
        "hl": hl,
        "start": start,
        "safe": safe,
        "gl": gl,
        "cr": cr,
        "lr": lr,
    }
    
    # 删除值为None的项
    params = {k: v for k, v in params.items() if v is not None}
    
    # 打印构建的URL（忠实显示所有参数）
    from urllib.parse import urlencode
    debug_url = f"https://www.google.com/search?{urlencode(params)}"
    page_num = (start // 10) + 1  # 计算页码（每页10个结果）
    print(f"  构建搜索URL (第{page_num}页): {debug_url}")
    
    resp = requests.get(
        url="https://www.google.com/search",
        headers={
            "User-Agent": get_useragent(),
            "Accept": "*/*"
        },
        params=params,
        proxies=proxies,
        timeout=timeout,
        verify=ssl_verify,
        cookies = {
            'CONSENT': 'PENDING+987', # Bypasses the consent page
            'SOCS': 'CAESHAgBEhIaAB',
        }
    )
    resp.raise_for_status()
    return resp


class SearchResult:
    def __init__(self, url, title, description):
        self.url = url
        self.title = title
        self.description = description

    def __repr__(self):
        return f"SearchResult(url={self.url}, title={self.title}, description={self.description})"


def search(q, num_results=10, hl=None, proxy=None, advanced=False, sleep_interval=0, timeout=5, safe="active", ssl_verify=None, gl=None, start_num=0, unique=False, cr=None, lr=None):
    """Search the Google search engine"""

    # Proxy setup
    proxies = {"https": proxy, "http": proxy} if proxy and (proxy.startswith("https") or proxy.startswith("http") or proxy.startswith("socks5")) else None

    start = start_num
    fetched_results = 0  # Keep track of the total fetched results
    fetched_links = set() # to keep track of links that are already seen previously

    while fetched_results < num_results:
        # Send request
        resp = _req(q, num_results - start,
                        start, proxies, timeout, safe, ssl_verify, gl, hl, cr, lr)
        
        # put in file - comment for debugging purpose
        # with open('google.html', 'w') as f:
        #     f.write(resp.text)
        
        # Parse
        soup = BeautifulSoup(resp.text, "html.parser")
        result_block = soup.find_all("div", class_="ezO2md")
        print(f"找到 {len(result_block)} 个 ezO2md 元素")
        new_results = 0  # Keep track of new results in this iteration

        for result in result_block:
            # Find the link tag within the result block
            link_tag = result.find("a", href=True)
            # Find the title tag within the link tag
            title_tag = link_tag.find("span", class_="CVA68e") if link_tag else None
            # Find the description tag within the result block
            description_tag = result.find("span", class_="FrIlee")

            # Extract and decode the link URL
            link = unquote(link_tag["href"].split("&")[0].replace("/url?q=", "")) if link_tag else ""
            
            # Check if all necessary tags are found
            if link_tag and title_tag and description_tag:
                print(f"    第{fetched_results+1}个链接: {link}，找到 {len(link_tag)} 个 link_tag 元素，")
            
            # 过滤无效链接：以/search?开头的无效搜索链接
            if link.startswith("/search?"):
                continue  # Skip this invalid link
            
            # Check if the link has already been fetched and if unique results are required
            if link in fetched_links and unique:
                continue  # Skip this result if the link is not unique
            # Add the link to the set of fetched links
            fetched_links.add(link)
            # Extract the title text
            title = title_tag.text if title_tag else ""
            # Extract the description text
            description = description_tag.text if description_tag else ""

            if link and title:
                # Increment the count of fetched results
                fetched_results += 1
                # Increment the count of new results in this iteration
                new_results += 1


            # Yield the result based on the advanced flag
            if advanced:
                yield SearchResult(link, title, description)  # Yield a SearchResult object
            else:
                yield link  # Yield only the link

            if fetched_results >= num_results:
                print(f"\n{'='*50}")  
                print(f"  已经达到本次搜索目标数量，停止搜索")
                print(f"{'='*50}")
                break  # Stop if we have fetched the desired number of results

        if new_results == 0:
            #If you want to have printed to your screen that the desired amount of queries can not been fulfilled, uncomment the line below:
            #print(f"Only {fetched_results} results found for query requiring {num_results} results. Moving on to the next query.")
            print(f"\n{'='*50}")  
            print(f"  已无更多搜索结果，停止搜索，本次搜索共找到 {fetched_results} 个结果")
            print(f"{'='*50}")
            break  # Break the loop if no new results were found in this iteration

        start += 10  # Prepare for the next set of results
        
        # 随机化请求间隔：sleep_interval ± 2秒范围内随机
        min_interval = max(0, sleep_interval - 2)  # 最小值不能为负数
        max_interval = sleep_interval + 2
        sleep_interval = random.uniform(min_interval, max_interval)
        print(f"  避免被Google封禁，本次休眠 {sleep_interval:.2f}秒")
        sleep(sleep_interval)

# ==================== googlesearch库代码结束 ====================

def detect_encoding_from_response_and_html(response, html_content=None):
    # 从响应头和HTML中提取编码信息
    try:
        # 第一步：从响应头中提取编码
        content_type = response.headers.get('Content-Type', '').lower()
        if 'charset=' in content_type:
            # 简化编码提取
            encoding = content_type.split('charset=')[1].split(';')[0].strip()
            return encoding
        
        # 从其他可能的头部字段中查找编码信息
        for header_name in ['charset', 'encoding']:
            if header_name in response.headers:
                return response.headers[header_name].strip()
        
        # 第二步：如果响应头没有编码信息，从HTML中查找
        if html_content:
            try:
                # 直接在HTML内容中查找 charset=
                import re
                # 匹配 <meta charset="..."> 和 <meta http-equiv="Content-Type" content="...charset=...">
                charset_match = re.search(r'charset\s*=\s*["\']?([^"\'>;\s]+)', html_content, re.IGNORECASE)
                if charset_match:
                    return charset_match.group(1).strip()
            except Exception:
                pass
        
        return None
    except Exception:
        return None

def smart_encoding_detection(html_content, response):
    # 智能编码检测：优先使用响应头编码，仅在必要时使用chardet检测
    try:
        # 优先使用响应头和HTML中的编码信息
        detected_encoding = detect_encoding_from_response_and_html(response, html_content)
        if detected_encoding:
            # 如果检测到的编码与当前使用的不同，尝试重新解码
            if detected_encoding.lower() != response.encoding.lower():
                try:
                    return response.content.decode(detected_encoding, errors='ignore')
                except (UnicodeDecodeError, LookupError):
                    pass
        
        # 使用chardet检测（置信度>0.7）
        import chardet
        detected = chardet.detect(response.content)
        if (detected.get('confidence', 0) > 0.7 and 
            detected.get('encoding', '').lower() != response.encoding.lower()):
            try:
                return response.content.decode(detected['encoding'], errors='ignore')
            except Exception:
                pass
        
        return html_content
    except Exception:
        return html_content

def fix_encoding_issues(html_content, response):
    # 修复网页内容的编码问题（保持向后兼容）
    return smart_encoding_detection(html_content, response)

def check_and_create_config_file():
    # 检查当前目录下是否存在search_config.xlsx，如果不存在则创建模板
    config_file = "search_config.xlsx"
    
    if os.path.exists(config_file):
        print(f"✓ 找到配置文件: {config_file}")
        return True
    else:
        print(f"✗ 未找到配置文件: {config_file}")
        print("正在创建配置模板...")
        create_config_template(config_file)
        return False

def create_config_template(filename):
    # 创建配置模板Excel文件，动态读取代码中的实际配置值
    try:
        # 动态读取代码中的实际配置值
        config_values = [
            GOOGLE_API_KEY,
            GOOGLE_SEARCH_ENGINE_ID, 
            SEARCH_KEYWORD,
            EXCLUDE_TERMS,
            SEARCH_METHOD,
            str(SEARCH_CONFIG['num_results']),
            SEARCH_CONFIG['language'],
            SEARCH_CONFIG['country'],
            str(OTHER_CONFIG['sleep_interval']),
            str(OTHER_CONFIG['sleep_interval']),
            str(CONTENT_EXTRACTION_CONFIG['extraction_workers']),
            str(CONTENT_EXTRACTION_CONFIG['enable_truncate']),
            str(CONTENT_EXTRACTION_CONFIG['max_chars']),
            str(TRANSLATION_CONFIG['enable_translation']),
            TRANSLATION_CONFIG['api_key'],
            TRANSLATION_CONFIG['api_url'],
            TRANSLATION_CONFIG['agent_id'],
            TRANSLATION_CONFIG['source_lang'],
            TRANSLATION_CONFIG['target_lang'],
            TRANSLATION_CONFIG['strategy'],
            str(TRANSLATION_CONFIG['translate_content']),
            str(TRANSLATION_CONFIG['translate_summary']),
            str(TRANSLATION_CONFIG['max_workers']),
            str(TRANSLATION_CONFIG['request_delay']),
            str(EMAIL_CONFIG['enable_email']),
        EMAIL_CONFIG['smtp_server'],
        str(EMAIL_CONFIG['smtp_port']),
        EMAIL_CONFIG['sender_email'],
        EMAIL_CONFIG['sender_auth_code'],
        EMAIL_CONFIG['recipient_email']
        ]
        
        # 创建配置模板数据，按组别分组并添加空行
        config_items = []
        config_chinese_names = []
        config_values_list = []
        config_descriptions = []
        config_required = []
        
        # 其他配置数据（正文提取配置组）
        other_config_items = []
        other_config_chinese_names = []
        other_config_values_list = []
        other_config_descriptions = []
        other_config_required = []
        
        # 1. Google搜索配置组（在"搜索关键词"上方插入空白行）
        # 写入前四项：搜索方法和API密钥配置
        config_items.extend([
            'SEARCH_METHOD',
            '',  # 空行分隔
            'SERP_API_KEY',
            '',  # 空行分隔
            'GOOGLE_API_KEY',
            'GOOGLE_SEARCH_ENGINE_ID'
        ])
        config_chinese_names.extend([
            '搜索接口',
            '',  # 空行分隔
            'SERP API密钥',
            '',  # 空行分隔
            'Google官方API密钥',
            'Google搜索引擎ID'
        ])
        config_values_list.extend([
            config_values[4],  # SEARCH_METHOD
            '',                # 空行分隔
            SERP_API_KEY,      # SERP_API_KEY
            '',                # 空行分隔
            config_values[0],  # GOOGLE_API_KEY
            config_values[1]   # GOOGLE_SEARCH_ENGINE_ID
        ])
        config_descriptions.extend([
            'google_api、googlesearch（目前暂时无法使用）或serp_api（第三方Google搜索API，在https://serpapi.com申请）',
            '',  # 空行分隔
            '使用SERP API时必填，在https://serpapi.com 注册获取',
            '',  # 空行分隔
            '使用Google官方API时必填，在https://programmablesearchengine.google.com/获取',
            '使用Google官方API时必填，在https://programmablesearchengine.google.com/ 获取'
        ])
        config_required.extend([
            '必填',
            '',  # 空行分隔
            '条件必填',
            '',  # 空行分隔
            '条件必填',
            '条件必填'
        ])

        # 插入一行空白（位于"搜索关键词"上方）
        config_items.append('')
        config_chinese_names.append('')
        config_values_list.append('')
        config_descriptions.append('')
        config_required.append('')

        # 再写入"搜索关键词"和"排除关键词"
        config_items.extend([
            'SEARCH_KEYWORD',
            'EXCLUDE_TERMS'
        ])
        config_chinese_names.extend([
            '搜索关键词',
            '排除关键词'
        ])
        config_values_list.extend([
            config_values[2],  # SEARCH_KEYWORD
            config_values[3]   # EXCLUDE_TERMS
        ])
        config_descriptions.extend([
            '用空格分隔多个搜索关键词',
            '用空格分隔多个排除关键词'
        ])
        config_required.extend([
            '必填',
            '*选填'
        ])
        
        # 添加空行
        config_items.append('')
        config_chinese_names.append('')
        config_values_list.append('')
        config_descriptions.append('')
        config_required.append('')
        
        # 2. 搜索参数配置组
        config_items.extend([
            'num_results',
            'language',
            'country'
        ])
        config_chinese_names.extend([
            '搜索结果数量',
            '搜索语言',
            '搜索国家'
        ])
        config_values_list.extend([
            config_values[5],  # num_results
            config_values[6],  # language
            config_values[7]   # country
        ])
        config_descriptions.extend([
            'google_api限制最多100，serp_api不限制（建议300）',
            'Google默认设置（不限制语言），auto（自动检测）、lang_en（英语）、lang_es（西班牙语）、lang_ar（阿拉伯语）、lang_fr（法语）、lang_pt（葡萄牙语）等',
            'Google默认设置（不限制国家），US（美国）、GB（英国）、FR（法国）、BR（巴西）、AR（阿根廷）、MX（墨西哥）等'
        ])
        config_required.extend([
            '必填',
            '*选填',
            '*选填'
        ])

        # 在"搜索国家"下新增：是否启用多关键词搜索模式
        config_items.append('enable_multi_keywords_mode')
        config_chinese_names.append('是否启用多关键词搜索模式')
        config_values_list.append('False')
        config_descriptions.append('True或False，启用"多搜索关键词模式"工作表进行批量搜索')
        config_required.append('必填')
        
        # 添加空行
        config_items.append('')
        config_chinese_names.append('')
        config_values_list.append('')
        config_descriptions.append('')
        config_required.append('')
        
        # 3. 其他配置组 - 移到其他配置工作表
        other_config_items.extend([
            'sleep_interval'
        ])
        other_config_chinese_names.extend([
            '搜索请求间隔'
        ])
        other_config_values_list.extend([
            config_values[8]  # sleep_interval
        ])
        other_config_descriptions.extend([
            '搜索请求间隔时间（秒），API调用间的延迟时间'
        ])
        other_config_required.extend([
            '*选填'
        ])
        
        # 添加空行
        config_items.append('')
        config_chinese_names.append('')
        config_values_list.append('')
        config_descriptions.append('')
        config_required.append('')
        
        # 4. 正文提取配置组 - 移到其他配置工作表
        other_config_items.extend([
            'enable_truncate',
            'extraction_workers',
            'max_chars'
        ])
        other_config_chinese_names.extend([
            '是否启用正文截取',
            '正文提取线程数',
            '正文截取最大字符数'
        ])
        other_config_values_list.extend([
            config_values[11],  # enable_truncate
            config_values[10],  # extraction_workers
            config_values[12]   # max_chars
        ])
        other_config_descriptions.extend([
            '是否启用正文截取：True或False（保持默认），超过指定字符数的正文将被截取，避免内容过长',
            '内容提取并发线程数（建议3-10个），保持默认',
            '最大字符数（启用截取时有效）'
        ])
        other_config_required.extend([
            '*选填',
            '*选填',
            '*选填'
        ])
        
        # 添加空行分隔符
        other_config_items.append('')
        other_config_chinese_names.append('')
        other_config_values_list.append('')
        other_config_descriptions.append('')
        other_config_required.append('')
        
        # 添加空行
        config_items.append('')
        config_chinese_names.append('')
        config_values_list.append('')
        config_descriptions.append('')
        config_required.append('')
        
        # 5. 翻译功能配置组
        config_items.extend([
            'enable_translation',
            'translation_api_key',
            'translation_api_url',
            'translation_agent_id',
            'translation_source_lang',
            'translation_target_lang',
            'translation_strategy',
            'translate_summary',
            'translate_content',
            'translation_max_workers',
            'translation_request_delay'
        ])
        config_chinese_names.extend([
            '是否启用翻译',
            '智谱翻译API密钥',
            '智谱翻译API地址',
            '翻译智能体ID',
            '源语言',
            '目标语言',
            '翻译策略',
            '是否翻译摘要',
            '是否翻译正文',
            '翻译线程数',
            '翻译请求延迟'
        ])
        config_values_list.extend([
            config_values[13],  # enable_translation
            config_values[14],  # translation_api_key
            config_values[15],  # translation_api_url
            config_values[16],  # translation_agent_id
            config_values[17],  # translation_source_lang
            config_values[18],  # translation_target_lang
            config_values[19],  # translation_strategy
            config_values[21],  # translate_summary
            config_values[20],  # translate_content
            config_values[22],  # translation_max_workers
            config_values[23]   # translation_request_delay
        ])
        config_descriptions.extend([
            '是否启用翻译功能：True或False',
            '启用翻译时必填，在bigmodel.cn获取',
            '保持默认',
            '保持默认',
            '源语言：auto（自动检测）',
            'zh-CN（简体中文）,wyw（文言文），en（英语），ja（日语），es（西班牙语），ar（阿拉伯语），de（德语），fr（法语），pt（葡萄牙语）',
            'two_step（两步翻译）效果最好，general（通用翻译）次之',
            'True或False',
            'True或False',
            '建议3-10个（数字越大速度越快）',
            'API请求延迟（秒），保持默认'
        ])
        config_required.extend([
            '*选填',
            '条件必填',
            '*选填',
            '*选填',
            '*选填',
            '*选填',
            '*选填',
            '*选填',
            '*选填',
            '*选填',
            '*选填'
        ])
        
        # 添加空行
        config_items.append('')
        config_chinese_names.append('')
        config_values_list.append('')
        config_descriptions.append('')
        config_required.append('')
        
        # 6. 邮件发送配置组
        config_items.extend([
            'enable_email',
            'smtp_server',
            'smtp_port',
            'sender_email',
            'sender_auth_code',
            'recipient_email'
        ])
        config_chinese_names.extend([
            '是否启用邮件',
            'SMTP服务器',
            'SMTP端口',
            '发件人邮箱',
            '发件人邮箱授权码',
            '收件人邮箱'
        ])
        config_values_list.extend([
            config_values[24],  # enable_email
            config_values[25],  # smtp_server
            config_values[26],  # smtp_port
            config_values[27],  # sender_email
            config_values[28],  # sender_auth_code
            config_values[29]   # recipient_email
        ])
        config_descriptions.extend([
            '是否启用邮件发送功能：True或False',
            'SMTP服务器地址',
            'SMTP端口号',
            '启用邮件时必填',
            '启用邮件时必填，在https://service.mail.qq.com/detail/0/75获取',
            '启用邮件时必填'
        ])
        config_required.extend([
            '*选填',
            '*选填',
            '*选填',
            '条件必填',
            '条件必填',
            '条件必填'
        ])
        
        # 创建配置模板数据（表头：配置项（请勿修改））
        config_data = {
            '配置项中文名': config_chinese_names,
            '配置项（请勿修改）': config_items,
            '配置值（请填入配置信息）': config_values_list,
            '必填/选填': config_required,
            '说明': config_descriptions
        }
        
        # 创建DataFrame
        df = pd.DataFrame(config_data)
        
        # 创建其他配置数据
        other_config_data = {
            '配置项中文名': other_config_chinese_names,
            '配置项': other_config_items,
            '配置值（请填入配置信息）': other_config_values_list,
            '必填/选填': other_config_required,
            '说明': other_config_descriptions
        }
        
        # 创建其他配置DataFrame
        other_df = pd.DataFrame(other_config_data)
        
        # 保存为Excel文件
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='配置模板', index=False)

            # 创建多搜索关键词模式工作表（空模板，供用户批量填写）
            # 预填一行说明 + 3 行示例，便于用户直接填写
            # 列顺序：搜索关键词、排除关键词、搜索结果数量、搜索语言、搜索国家、说明
            multi_keywords_rows = [
                {
                    '搜索关键词': 'air purifiers wholesaler',
                    '排除关键词': '',
                    '搜索结果数量': 10,
                    '搜索语言': '',
                    '搜索国家': '',
                    '说明': '可以多行填写"关键词"和"搜索结果数量"，在一次运行中，会自动搜索多条关键词。'
                },
                {
                    '搜索关键词': 'solar panels distributors',
                    '排除关键词': '',
                    '搜索结果数量': 10,
                    '搜索语言': '',
                    '搜索国家': '',
                    '说明': ''
                },
                {
                    '搜索关键词': 'hospital supplies importer',
                    '排除关键词': '',
                    '搜索结果数量': 10,
                    '搜索语言': '',
                    '搜索国家': '',
                    '说明': ''
                }
            ]
            multi_keywords_df = pd.DataFrame(multi_keywords_rows, columns=['搜索关键词','排除关键词','搜索结果数量','搜索语言','搜索国家','说明'])
            # 第二张工作表
            multi_keywords_df.to_excel(writer, sheet_name='多搜索关键词模式', index=False)

            # 第三张：其他配置
            other_df.to_excel(writer, sheet_name='其他配置', index=False)
            
            # 创建搜索语种工作表
            # 创建搜索语言代码工作表数据
            language_mapping = [
                ('Google默认', 'Google默认'),
                ('auto', '自动检测'),
                ('lang_en', '英语'),
                ('lang_zh-CN', '简体中文'),
                ('lang_zh-TW', '繁体中文'),
                ('lang_es', '西班牙语'),
                ('lang_ja', '日语'),
                ('lang_ko', '韩语'),
                ('lang_ru', '俄语'),
                ('lang_vi', '越南语'),
                ('lang_ar', '阿拉伯语'),
                ('lang_de', '德语'),
                ('lang_fr', '法语'),
                ('lang_pt', '葡萄牙语')
            ]
            
            language_codes = [item[0] for item in language_mapping]
            language_names = [item[1] for item in language_mapping]
            language_note = '更多语言代码，请查阅https://developers.google.com/custom-search/v1/reference/rest/v1/cse/list?hl=zh-cn'
            language_desc = [language_note] + [''] * (len(language_codes) - 1)
            
            language_data = {
                '语种代码': language_codes,
                '语种名称': language_names,
                '说明': language_desc
            }
            language_df = pd.DataFrame(language_data)
            language_df.to_excel(writer, sheet_name='搜索语言代码', index=False)
            
            # 创建搜索国家工作表（代码表）
            country_codes = ['Google默认', 'CN', 'US', 'GB', 'JP', 'KR', 'IN', 'FR', 'DE', 'RU', 'ES', 'BR', 'AR', 'MX', 'VN']
            country_names = ['Google默认', '中国', '美国', '英国', '日本', '韩国', '印度', '法国', '德国', '俄罗斯', '西班牙', '巴西', '阿根廷', '墨西哥', '越南']
            country_note = '更多国家代码请查看 https://www.alibabacloud.com/help/zh/chatapp/country-or-region-codes'
            country_desc = [country_note] + [''] * (len(country_codes) - 1)
            country_data = {
                '国家代码': country_codes,
                '国家名称': country_names,
                '说明': country_desc
            }
            country_df = pd.DataFrame(country_data)
            country_df.to_excel(writer, sheet_name='搜索国家代码', index=False)
            
            # 获取工作表对象以调整列宽和字体
            worksheet = writer.sheets['配置模板']
            other_worksheet = writer.sheets['其他配置']
            language_worksheet = writer.sheets['搜索语言代码']
            country_worksheet = writer.sheets['搜索国家代码']
            multi_keywords_worksheet = writer.sheets['多搜索关键词模式']
            
            # 导入字体样式和数据验证
            from openpyxl.styles import Font, PatternFill, Alignment
            from openpyxl.worksheet.datavalidation import DataValidation
            
            # 设置等线字体样式
            default_font = Font(name='等线', size=11)
            header_font = Font(name='等线', size=11, bold=True)
            
            # 设置表头样式
            header_fill = PatternFill(start_color='E6F3FF', end_color='E6F3FF', fill_type='solid')
            
            # 设置必填项高亮样式
            required_fill = PatternFill(start_color='FFE6CC', end_color='FFE6CC', fill_type='solid')  # 浅橙色背景
            required_font = Font(name='等线', size=11, bold=True, color='CC0000')  # 红色加粗字体
            
            # 应用字体到所有单元格的函数
            def apply_styles_to_worksheet(ws):
                for row in ws.iter_rows():
                    if row[0].row == 1:  # 表头行
                        for cell in row:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                    else:  # 数据行
                        # 首先检查是否为功能开关行（搜索接口、是否启用翻译、是否启用邮件）
                        is_function_switch = False
                        if len(row) > 1 and row[1].value in ['SEARCH_METHOD', 'enable_translation', 'enable_email']:  # 第2列是"配置项"列
                            is_function_switch = True
                        
                        if is_function_switch:
                            # 整行应用黑底白字样式
                            black_fill = PatternFill(start_color='000000', end_color='000000', fill_type='solid')
                            white_font = Font(name='等线', size=11, bold=True, color='FFFFFF')
                            for col_cell in row:
                                col_cell.font = white_font
                                col_cell.fill = black_fill
                                col_cell.alignment = Alignment(vertical='top', wrap_text=False)
                        else:
                            # 其他行的样式处理
                            # 检查是否为需要橙色背景的配置项
                            config_item_value = None
                            if len(row) > 1:  # 确保有第2列（配置项列）
                                config_item_value = row[1].value
                            
                            # 需要橙色背景的配置项列表
                            orange_bg_items = ['SEARCH_KEYWORD', 'SERP_API_KEY', 'GOOGLE_API_KEY', 'GOOGLE_SEARCH_ENGINE_ID']
                            
                            for cell in row:
                                # 检查是否为必填项或特定的API密钥配置项
                                if (cell.column == 4 and cell.value == '必填') or (config_item_value in orange_bg_items):
                                    # 整行应用必填项样式（橙色背景，使用默认字体）
                                    for col_cell in row:
                                        col_cell.font = default_font
                                        col_cell.fill = required_fill
                                        col_cell.alignment = Alignment(vertical='top', wrap_text=False)
                                    break
                                elif cell.column == 4 and cell.value == '条件必填' and config_item_value not in orange_bg_items:
                                    # 其他条件必填项样式（去除背景色，使用默认字体）
                                    for col_cell in row:
                                        col_cell.font = default_font
                                        col_cell.fill = PatternFill()  # 默认背景色（无填充）
                                        col_cell.alignment = Alignment(vertical='top', wrap_text=False)
                                    break
                            else:
                                # 普通数据行
                                for cell in row:
                                    cell.font = default_font
                                    cell.alignment = Alignment(vertical='top', wrap_text=False)
            
            # 为所有工作表应用样式
            apply_styles_to_worksheet(worksheet)
            apply_styles_to_worksheet(other_worksheet)
            apply_styles_to_worksheet(language_worksheet)
            apply_styles_to_worksheet(country_worksheet)
            apply_styles_to_worksheet(multi_keywords_worksheet)
            
            # 为各种配置项添加下拉列表
            def add_data_validation(ws):
                # 预创建并复用数据验证对象，避免在循环中重复创建
                true_false_validation = DataValidation(type="list", formula1='"True,False"', allow_blank=False)
                search_method_validation = DataValidation(type="list", formula1='"google_api,googlesearch,serp_api"', allow_blank=False)
                language_validation = DataValidation(type="list", formula1="搜索语言代码!$A$2:$A$100", allow_blank=False)
                country_validation = DataValidation(type="list", formula1="搜索国家代码!$A$2:$A$100", allow_blank=False)
                strategy_validation = DataValidation(type="list", formula1='"two_step,general"', allow_blank=False)

                # 将验证对象添加到工作表（先添加，再批量绑定单元格）
                ws.add_data_validation(true_false_validation)
                ws.add_data_validation(search_method_validation)
                ws.add_data_validation(language_validation)
                ws.add_data_validation(country_validation)
                ws.add_data_validation(strategy_validation)

                # 遍历行，将对应单元格加入相应的验证
                for row_num, row in enumerate(ws.iter_rows(), 1):
                    if len(row) > 1 and row[1].value in ['enable_translation', 'enable_email', 'enable_truncate', 'translate_summary', 'translate_content', 'enable_multi_keywords_mode']:
                        true_false_validation.add(ws.cell(row=row_num, column=3))
                    elif len(row) > 1 and row[1].value == 'SEARCH_METHOD':
                        search_method_validation.add(ws.cell(row=row_num, column=3))
                    elif len(row) > 1 and row[1].value == 'language':
                        language_validation.add(ws.cell(row=row_num, column=3))
                    elif len(row) > 1 and row[1].value == 'country':
                        country_validation.add(ws.cell(row=row_num, column=3))
                    elif len(row) > 1 and row[1].value == 'translation_strategy':
                        strategy_validation.add(ws.cell(row=row_num, column=3))
            
            # 为配置模板工作表添加数据验证
            add_data_validation(worksheet)
            
            # 通用下拉验证复用函数
            def add_list_validation_range(target_ws, column_letter, start_row, end_row, source_range, allow_blank=True):
                try:
                    validation = DataValidation(type="list", formula1=source_range, allow_blank=allow_blank)
                    validation.add(f"{column_letter}{start_row}:{column_letter}{end_row}")
                    target_ws.add_data_validation(validation)
                except Exception:
                    # 忽略非关键错误
                    pass

            # 为"多搜索关键词模式"添加下拉：语言(D)与国家(E)
            add_list_validation_range(multi_keywords_worksheet, 'D', 2, 1000, "搜索语言代码!$A$2:$A$100", allow_blank=True)
            add_list_validation_range(multi_keywords_worksheet, 'E', 2, 1000, "搜索国家代码!$A$2:$A$100", allow_blank=True)
            
            # 自动调整列宽的函数
            def adjust_column_width(ws):
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if cell.value is not None:
                                # 计算中文字符长度（中文字符按2个字符计算）
                                cell_value = str(cell.value)
                                chinese_chars = len([c for c in cell_value if '\u4e00' <= c <= '\u9fff'])
                                ascii_chars = len(cell_value) - chinese_chars
                                cell_length = ascii_chars + chinese_chars * 2
                                if cell_length > max_length:
                                    max_length = cell_length
                        except:
                            pass
                    
                    # 设置合适的列宽，最小15，最大100
                    adjusted_width = max(15, min(max_length + 3, 100))
                    ws.column_dimensions[column_letter].width = adjusted_width
            
            # 为所有工作表调整列宽
            adjust_column_width(worksheet)
            adjust_column_width(other_worksheet)
            adjust_column_width(language_worksheet)
            adjust_column_width(country_worksheet)
            adjust_column_width(multi_keywords_worksheet)
        
        print(f"✓ 配置模板已创建: {filename}")
        
        return True
        
    except Exception as e:
        print(f"✗ 创建配置模板失败: {e}")
        return False

def load_config_from_excel(filename="search_config.xlsx"):
    # 从Excel文件加载配置
    try:
        if not os.path.exists(filename):
            print(f"配置文件不存在: {filename}")
            return False
            
        # 读取Excel文件
        df = pd.read_excel(filename, sheet_name='配置模板')
        # 兼容列名：将"配置项（请勿修改）"统一为"配置项"用于内部处理
        if '配置项（请勿修改）' in df.columns and '配置项' not in df.columns:
            df = df.rename(columns={'配置项（请勿修改）': '配置项'})
        
        # 读取其他配置工作表
        try:
            other_df = pd.read_excel(filename, sheet_name='其他配置')
            if '配置项（请勿修改）' in other_df.columns and '配置项' not in other_df.columns:
                other_df = other_df.rename(columns={'配置项（请勿修改）': '配置项'})
            # 合并两个工作表的数据
            df = pd.concat([df, other_df], ignore_index=True)
        except:
            # 如果其他配置工作表不存在，继续使用主工作表
            pass
        
        # 创建配置字典
        config_dict = {}
        for _, row in df.iterrows():
            key = row['配置项']
            value = row['配置值（请填入配置信息）']
            
            # 处理布尔值
            if str(value).lower() in ['true', 'false']:
                value = str(value).lower() == 'true'
            # 处理数字
            elif str(value).isdigit():
                value = int(value)
            # 处理浮点数
            elif str(value).replace('.', '').isdigit():
                value = float(value)
            # 处理空值
            elif pd.isna(value) or str(value).strip() == '':
                value = ''
            
            config_dict[key] = value
        
        # 更新全局配置
        global GOOGLE_API_KEY, GOOGLE_SEARCH_ENGINE_ID, SERP_API_KEY, SEARCH_KEYWORD, EXCLUDE_TERMS, SEARCH_METHOD
        global SEARCH_CONFIG, CONTENT_EXTRACTION_CONFIG, TRANSLATION_CONFIG, EMAIL_CONFIG
        
        # 更新基本配置
        # 分别读取GOOGLE_API_KEY和SERP_API_KEY
        GOOGLE_API_KEY = config_dict.get('GOOGLE_API_KEY', '')
        SERP_API_KEY = config_dict.get('SERP_API_KEY', '')
        GOOGLE_SEARCH_ENGINE_ID = config_dict.get('GOOGLE_SEARCH_ENGINE_ID', GOOGLE_SEARCH_ENGINE_ID)
        SEARCH_KEYWORD = config_dict.get('SEARCH_KEYWORD', SEARCH_KEYWORD)
        EXCLUDE_TERMS = config_dict.get('EXCLUDE_TERMS', EXCLUDE_TERMS)
        SEARCH_METHOD = config_dict.get('SEARCH_METHOD', SEARCH_METHOD)
        
        # 更新搜索配置
        SEARCH_CONFIG['num_results'] = config_dict.get('num_results', SEARCH_CONFIG['num_results'])
        # 处理'Google默认'配置：将其转换为空字符串表示使用默认设置
        language_value = config_dict.get('language', SEARCH_CONFIG['language'])
        SEARCH_CONFIG['language'] = '' if language_value == 'Google默认' else language_value
        country_value = config_dict.get('country', SEARCH_CONFIG['country'])
        SEARCH_CONFIG['country'] = '' if country_value == 'Google默认' else country_value

        # 读取是否启用多关键词模式（如存在则覆盖运行时总开关）
        global ENABLE_MULTI_KEYWORDS_MODE
        if 'enable_multi_keywords_mode' in config_dict:
            ENABLE_MULTI_KEYWORDS_MODE = bool(config_dict.get('enable_multi_keywords_mode'))
        
        # 更新其他配置
        OTHER_CONFIG['sleep_interval'] = config_dict.get('sleep_interval', 1)
        
        # 更新正文提取配置
        CONTENT_EXTRACTION_CONFIG['extraction_workers'] = config_dict.get('extraction_workers', CONTENT_EXTRACTION_CONFIG['extraction_workers'])
        CONTENT_EXTRACTION_CONFIG['enable_truncate'] = config_dict.get('enable_truncate', CONTENT_EXTRACTION_CONFIG['enable_truncate'])
        CONTENT_EXTRACTION_CONFIG['max_chars'] = config_dict.get('max_chars', CONTENT_EXTRACTION_CONFIG['max_chars'])
        
        # 更新翻译配置
        TRANSLATION_CONFIG['enable_translation'] = config_dict.get('enable_translation', TRANSLATION_CONFIG['enable_translation'])
        TRANSLATION_CONFIG['api_key'] = config_dict.get('translation_api_key', TRANSLATION_CONFIG['api_key'])
        TRANSLATION_CONFIG['api_url'] = config_dict.get('translation_api_url', TRANSLATION_CONFIG['api_url'])
        TRANSLATION_CONFIG['agent_id'] = config_dict.get('translation_agent_id', TRANSLATION_CONFIG['agent_id'])
        TRANSLATION_CONFIG['source_lang'] = config_dict.get('translation_source_lang', TRANSLATION_CONFIG['source_lang'])
        TRANSLATION_CONFIG['target_lang'] = config_dict.get('translation_target_lang', TRANSLATION_CONFIG['target_lang'])
        TRANSLATION_CONFIG['strategy'] = config_dict.get('translation_strategy', TRANSLATION_CONFIG['strategy'])
        TRANSLATION_CONFIG['translate_content'] = config_dict.get('translate_content', TRANSLATION_CONFIG['translate_content'])
        TRANSLATION_CONFIG['translate_summary'] = config_dict.get('translate_summary', TRANSLATION_CONFIG['translate_summary'])
        TRANSLATION_CONFIG['max_workers'] = config_dict.get('translation_max_workers', TRANSLATION_CONFIG['max_workers'])
        TRANSLATION_CONFIG['request_delay'] = config_dict.get('translation_request_delay', TRANSLATION_CONFIG['request_delay'])
        
        # 更新邮件配置
        EMAIL_CONFIG['enable_email'] = config_dict.get('enable_email', EMAIL_CONFIG['enable_email'])
        EMAIL_CONFIG['smtp_server'] = config_dict.get('smtp_server', EMAIL_CONFIG['smtp_server'])
        EMAIL_CONFIG['smtp_port'] = config_dict.get('smtp_port', EMAIL_CONFIG['smtp_port'])
        EMAIL_CONFIG['sender_email'] = config_dict.get('sender_email', EMAIL_CONFIG['sender_email'])
        EMAIL_CONFIG['sender_auth_code'] = config_dict.get('sender_auth_code', EMAIL_CONFIG['sender_auth_code'])
        EMAIL_CONFIG['recipient_email'] = config_dict.get('recipient_email', EMAIL_CONFIG['recipient_email'])
        
        return True
        
    except Exception as e:
        print(f"✗ 从Excel文件加载配置失败: {e}")
        return False

def format_number(num):
    # 将大数字格式化为易读形式
    if num >= 10000:
        return f"{num / 10000:.1f}万".rstrip('0').rstrip('.')
    else:
        return str(num)

def detect_search_language(keyword):
    # 检测搜索关键词的语言
    if not LANGDETECT_AVAILABLE:
        return 'en'  # 默认英语
    
    try:
        # 检测语言（标准化为小写并用连字符）
        detected_lang = detect(keyword)  # 如 'es', 'en', 'zh-cn'
        detected_lang = detected_lang.lower().replace('_', '-')
        return detected_lang or 'en'
        
    except Exception as e:
        return 'en'

# 统一语言代码显示名称映射模块
def get_language_display_name_from_excel(lang_code: str, filename: str = "search_config.xlsx") -> str:
    """
    从"搜索语言代码"工作表中查找对应的语种名称
    """
    if not lang_code:
        return "Google默认"
    
    try:
        # 读取搜索语言代码工作表
        df = pd.read_excel(filename, sheet_name='搜索语言代码')
        
        # 查找对应的语种名称
        matching_row = df[df['语种代码'] == lang_code]
        if not matching_row.empty:
            return matching_row.iloc[0]['语种名称']
        else:
            # 如果找不到，返回代码本身
            return lang_code
    except Exception:
        # 如果读取失败，回退到原来的方法
        return get_language_display_name(lang_code)

def get_country_display_name_from_excel(country_code: str, filename: str = "search_config.xlsx") -> str:
    """
    从"搜索国家代码"工作表中查找对应的国家名称
    """
    if not country_code:
        return "Google默认"
    
    try:
        # 读取搜索国家代码工作表
        df = pd.read_excel(filename, sheet_name='搜索国家代码')
        
        # 查找对应的国家名称
        matching_row = df[df['国家代码'] == country_code]
        if not matching_row.empty:
            return matching_row.iloc[0]['国家名称']
        else:
            # 如果找不到，返回代码本身
            return country_code
    except Exception:
        # 如果读取失败，返回原始代码
        return country_code

def get_language_display_name(lang_code: str) -> str:
    """
    将语言代码转换为中文显示名称。
    支持形如 'en', 'es', 'zh-CN', 'zh-TW'，以及带 'lang_' 前缀的代码。
    未知代码返回其大写形式。
    """
    if not lang_code or lang_code == 'Google默认':
        return "Google默认设置（不限制语言）"
    base_code = lang_code.replace('lang_', '') if lang_code.startswith('lang_') else lang_code
    base_code_lower = base_code.lower()
    return LANGUAGE_DISPLAY_MAP.get(base_code_lower, base_code.upper())

def get_unified_search_info_display(keywords_list=None, languages_list=None, countries_list=None):
    """
    获取统一的搜索信息显示，支持多关键词、多语种和多国家拼接
    用于HTML、DOCX等文件的搜索信息显示，与Excel保持一致
    
    Args:
        keywords_list: 关键词列表，如果为None则使用全局SEARCH_KEYWORD
        languages_list: 语言列表，如果为None则使用配置中的语言
        countries_list: 国家列表，如果为None则使用配置中的国家
    
    Returns:
        tuple: (keywords_display, languages_display, countries_display)
    """
    # 处理关键词显示
    if keywords_list is None:
        keywords_display = SEARCH_KEYWORD
    else:
        keywords_display = " | ".join(keywords_list) if len(keywords_list) > 1 else keywords_list[0]
    
    # 处理语言显示
    if languages_list is None:
        # 使用配置中的语言（与Excel逻辑一致）
        search_lang, display_name = get_search_language_code('google')
        if not search_lang:
            languages_display = get_language_display_name_from_excel('Google默认')
        else:
            languages_display = get_language_display_name_from_excel(search_lang)
    else:
        # 多语种拼接
        display_names = []
        for lang in languages_list:
            if lang:
                display_name = get_language_display_name_from_excel(lang)
                display_names.append(display_name)
        languages_display = " | ".join(display_names) if display_names else "Google默认"
    
    # 处理国家显示
    if countries_list is None:
        # 使用配置中的国家
        if not SEARCH_CONFIG['country'] or SEARCH_CONFIG['country'] == 'Google默认':
            countries_display = "Google默认"
        else:
            countries_display = get_country_display_name_from_excel(SEARCH_CONFIG['country'])
    else:
        # 多国家拼接
        display_names = []
        for country in countries_list:
            if country:
                display_name = get_country_display_name_from_excel(country)
                display_names.append(display_name)
        countries_display = " | ".join(display_names) if display_names else "Google默认"
    
    return keywords_display, languages_display, countries_display


def get_search_language_code(api_type='google'):
    """
    获取搜索语言代码，支持Google官方API和Serp API
    
    Args:
        api_type: 'google' 或 'serp'，指定API类型
    
    Returns:
        tuple: (language_code, display_name)
    """
    if SEARCH_CONFIG['language'] == 'auto':
        detected_lang = detect_search_language(SEARCH_KEYWORD)
        # 如果有特殊映射则使用，否则直接加 lang_ 前缀
        search_lang = SPECIAL_LANG_PREFIX_MAP.get(detected_lang, f'lang_{detected_lang}')
    elif SEARCH_CONFIG['language'] == '' or SEARCH_CONFIG['language'] == 'Google默认':
        # 空字符串或'Google默认'表示使用Google默认设置（不限制语言）
        search_lang = ''
    else:
        # 手动指定的语言代码（用户已包含 lang_ 前缀）
        search_lang = SEARCH_CONFIG['language']
    
    # 生成显示名称（中文友好名）
    display_name = get_language_display_name(search_lang)
    
    return search_lang, display_name

# 创建线程锁，用于控制并发请求
translation_lock = threading.Lock()

# 全局变量用于跟踪当前状态，支持中断保存
current_search_results = []
current_extracted_results = []
current_translated_results = []
interrupted = False

# 全局网页计数器，跟踪整个运行过程中的网页数量
global_webpage_counter = 0

def signal_handler(signum, frame):
    # 处理用户中断信号（Ctrl+C）
    global interrupted
    print("\n\n" + "="*60)
    print("⚠ 检测到用户中断信号 (Ctrl+C)")
    print("="*60)
    print("正在保存已获取的数据...")
    interrupted = True
    
    # 保存当前已获取的数据
    save_current_results()
    
    print("="*60)
    print("数据保存完成！程序即将退出。")
    print("="*60)
    sys.exit(0)

def emergency_exit_handler():
    # 紧急退出处理器（用于atexit注册）
    # 处理程序正常退出、系统关机、用户注销等情况
    global interrupted
    if not interrupted:  # 避免重复保存
        print("\n⚠ 检测到程序意外退出，正在紧急保存数据...")
        save_current_results()
        print("✓ 紧急数据保存完成")

# Windows控制台事件处理
if sys.platform == "win32":
    try:
        import ctypes
        from ctypes import wintypes
        
        # Windows控制台事件常量
        CTRL_C_EVENT = 0
        CTRL_BREAK_EVENT = 1
        CTRL_CLOSE_EVENT = 2
        CTRL_LOGOFF_EVENT = 5
        CTRL_SHUTDOWN_EVENT = 6
        
        def windows_console_handler(event):
            """
            Windows控制台事件处理器
            处理控制台关闭、系统关机、用户注销等事件
            """
            global interrupted
            event_names = {
                CTRL_C_EVENT: "Ctrl+C",
                CTRL_BREAK_EVENT: "Ctrl+Break", 
                CTRL_CLOSE_EVENT: "控制台关闭",
                CTRL_LOGOFF_EVENT: "用户注销",
                CTRL_SHUTDOWN_EVENT: "系统关机"
            }
            
            event_name = event_names.get(event, f"未知事件({event})")
            print(f"\n⚠ 检测到Windows控制台事件: {event_name}")
            print("正在紧急保存数据...")
            
            interrupted = True
            save_current_results()
            print("✓ 数据保存完成")
            
            # 对于Ctrl+C和Ctrl+Break事件，强制退出程序
            if event in [CTRL_C_EVENT, CTRL_BREAK_EVENT]:
                print("程序即将退出...")
                sys.exit(0)
            
            return True  # 表示已处理该事件
        
        # 注册Windows控制台事件处理器
        kernel32 = ctypes.windll.kernel32
        handler_routine = ctypes.WINFUNCTYPE(wintypes.BOOL, wintypes.DWORD)(windows_console_handler)
        kernel32.SetConsoleCtrlHandler(handler_routine, True)
        
    except ImportError:
        print("警告: 无法导入Windows API，部分意外退出保护功能不可用")
    except Exception as e:
        print(f"警告: Windows控制台事件处理器注册失败: {e}")

def save_current_results():
    # 保存当前已获取的结果数据
    global current_search_results, current_extracted_results, current_translated_results
    
    if not current_search_results and not current_extracted_results and not current_translated_results:
        print("没有数据需要保存")
        return
    
    # 根据模式生成文件名：多关键词模式使用固定前缀，单关键词模式使用当前关键词
    if ENABLE_MULTI_KEYWORDS_MODE:
        base_filename = 'mutiple_keywords'
    else:
        base_filename = "".join(c for c in SEARCH_KEYWORD if c.isalnum() or c in (' ', '-', '_')).rstrip()
        base_filename = base_filename.replace(' ', '_')
    
    # 确定要保存的结果
    results_to_save = []
    if current_translated_results:
        results_to_save = current_translated_results
        print(f"保存翻译后的结果: {len(results_to_save)} 个")
    elif current_extracted_results:
        results_to_save = current_extracted_results
        print(f"保存提取后的结果: {len(results_to_save)} 个")
    elif current_search_results:
        results_to_save = current_search_results
        print(f"保存搜索结果: {len(results_to_save)} 个")
    
    if not results_to_save:
        print("没有有效数据需要保存")
        return
    
    # 生成带时间戳的文件名（精确到分钟）
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    docx_filename = f"{base_filename}_interrupted_{timestamp}.docx"
    html_filename = f"{base_filename}_interrupted_{timestamp}.html"
    xlsx_filename = f"{base_filename}_interrupted_{timestamp}.xlsx"
    
    try:
        # 保存为Word文档
        save_results_to_docx(results_to_save, docx_filename, requested_count=SEARCH_CONFIG['num_results'])
        print(f"✓ Word文档已保存: {docx_filename}")
    except Exception as e:
        print(f"✗ Word文档保存失败: {e}")
    
    try:
        # 保存为HTML文件
        save_results_to_html(results_to_save, html_filename, requested_count=SEARCH_CONFIG['num_results'])
        print(f"✓ HTML文件已保存: {html_filename}")
    except Exception as e:
        print(f"✗ HTML文件保存失败: {e}")
    
    try:
        # 保存为Excel文件
        save_results_to_excel(results_to_save, xlsx_filename)
        print(f"✓ Excel文件已保存: {xlsx_filename}")
    except Exception as e:
        print(f"✗ Excel文件保存失败: {e}")

def translate_text(text, text_type="content"):
    # 使用智谱翻译API翻译文本（支持并发）
    if not TRANSLATION_CONFIG['enable_translation'] or not text.strip():
        return text
    
    # 检查是否需要翻译
    if text_type == "content" and not TRANSLATION_CONFIG['translate_content']:
        return text
    if text_type == "summary" and not TRANSLATION_CONFIG['translate_summary']:
        return text
    
    try:
        # 添加请求延迟，避免API限流
        time.sleep(TRANSLATION_CONFIG['request_delay'])
        
        # 直接使用传入的文本进行翻译（文本已在调用前截取）
        text_to_translate = text
        
        payload = {
            "agent_id": TRANSLATION_CONFIG['agent_id'],
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": text_to_translate
                        }
                    ]
                }
            ],
            "custom_variables": {
                "source_lang": TRANSLATION_CONFIG['source_lang'],
                "target_lang": TRANSLATION_CONFIG['target_lang'],
                "strategy": TRANSLATION_CONFIG['strategy']
            }
        }
        
        headers = {
            "Authorization": f"Bearer {TRANSLATION_CONFIG['api_key']}",
            "Content-Type": "application/json"
        }
        
        with translation_lock:
            print(f"  正在翻译{text_type}...")
        
        response = requests.post(TRANSLATION_CONFIG['api_url'], json=payload, headers=headers, timeout=30)
        response.raise_for_status()
        
        result = response.json()
        
        if 'choices' in result and len(result['choices']) > 0:
            content = result['choices'][0]['messages'][0]['content']
            # 根据API文档，content是对象，包含type和text字段
            if isinstance(content, dict) and 'text' in content:
                translated_text = content['text']
                with translation_lock:
                    print(f"  {text_type}翻译完成")
                return translated_text
            elif isinstance(content, str):
                # 兼容直接返回字符串的情况
                translated_text = content
                with translation_lock:
                    print(f"  {text_type}翻译完成")
                return translated_text
            else:
                with translation_lock:
                    print(f"  {text_type}翻译失败：content格式异常")
                return text
        else:
            with translation_lock:
                print(f"  {text_type}翻译失败：API返回格式异常")
            return text
            
    except requests.exceptions.Timeout:
        with translation_lock:
            print(f"  {text_type}翻译超时，返回原文")
        return text
    except requests.exceptions.RequestException as e:
        with translation_lock:
            print(f"  {text_type}翻译请求失败：{e}")
        return text
    except Exception as e:
        with translation_lock:
            print(f"  {text_type}翻译异常：{e}")
        return text

def extract_single_article(task_data):
    # 提取单个网页内容（用于并发处理）
    import threading
    import time
    
    global_index, url, google_title, google_description = task_data
    
    print(f"  开始提取第 {global_index} 个网页: {url}")
    
    # 使用线程超时机制（Windows兼容）
    result_container = [None]
    exception_container = [None]
    
    def extract_worker():
        try:
            title, text, description, publish_date = extract_article_title_and_text(url, google_title, google_description)
            result_container[0] = (url, title, text, description, publish_date)
        except Exception as e:
            exception_container[0] = e
    
    # 启动提取线程
    thread = threading.Thread(target=extract_worker)
    thread.daemon = True
    thread.start()
    
    # 等待30秒或线程完成
    thread.join(timeout=30)
    
    if thread.is_alive():
        print(f"  ⚠ 第 {global_index} 个网页提取超时: {url}")
        return (url, google_title, "", google_description)
    elif exception_container[0]:
        print(f"  ✗ 第 {global_index} 个网页提取失败: {url}, 错误: {exception_container[0]}")
        return (url, google_title, "", google_description)
    else:
        return result_container[0]

def extract_content_phase(search_results, total_target_count=None, start_index=0):
    """
    第一阶段：并发提取所有搜索结果的内容（不翻译）
    
    Args:
        search_results: 搜索结果列表，格式为 [(url, title, description), ...]
        total_target_count: 用户指定的总搜索数量，用于显示进度
        start_index: 多关键词模式下的起始索引（用于进度显示）
    
    Returns:
        list: 提取结果列表，格式为 [(url, title, text, description, publish_date), ...]
    """
    global interrupted, current_extracted_results, global_webpage_counter
    
    print(f"\n{'='*50}")
    print("第一阶段：并发提取搜索结果内容")
    print(f"{'='*50}")
    
    # 准备提取任务，使用全局计数器
    extraction_tasks = []
    for i, (url, google_title, google_description) in enumerate(search_results):
        global_webpage_counter += 1
        extraction_tasks.append((global_webpage_counter, url, google_title, google_description))
    
    print(f"开始并发提取 {len(extraction_tasks)} 个网页内容...")
    print(f"并发线程数: {CONTENT_EXTRACTION_CONFIG['extraction_workers']}")  # 提取内容使用较少的线程数，避免对目标网站造成压力
    
    results = []
    
    # 使用线程池进行并发提取
    with concurrent.futures.ThreadPoolExecutor(max_workers=CONTENT_EXTRACTION_CONFIG['extraction_workers']) as executor:
        # 提交所有提取任务
        future_to_task = {executor.submit(extract_single_article, task): task for task in extraction_tasks}
        
        completed_count = 0
        total_count = len(extraction_tasks)
        
        # 收集提取结果
        for future in concurrent.futures.as_completed(future_to_task):
            # 检查是否被中断
            if interrupted:
                print("\n⚠ 检测到中断信号，停止内容提取...")
                break
                
            task = future_to_task[future]
            completed_count += 1
            
            try:
                result = future.result()
                results.append(result)
                
                # 实时更新全局变量，支持中断保存
                current_extracted_results = results.copy()
                
                # 显示进度：统一使用全局进度
                if total_target_count:
                    print(f"✓ 第 {task[0]} 个网页提取完成 ({task[0]}/{total_target_count})")
                else:
                    print(f"✓ 第 {task[0]} 个网页提取完成 ({task[0]}/{total_count})")
                
                # 立即显示该网页的详细信息
                if len(result) == 5:  # 包含发布时间
                    url, title, text, description, publish_date = result
                else:  # 向后兼容
                    url, title, text, description = result
                    publish_date = ""
                print(f"\n=== 第 {task[0]} 个搜索结果 ===")
                print(f"URL: {url}")
                print(f"最终标题: {title}")
                print(f"最终摘要: {description}")
                if publish_date:
                    print(f"发布时间: {publish_date}")
                print(f"正文内容（智能截取500字）：")
                if text:
                    # 截取前500个字符
                    truncated_text = text[:500] + "..." if len(text) > 500 else text
                    print(truncated_text)
                else:
                    print("（无正文内容）")
                print("=" * 50)
                
            except Exception as e:
                # 显示失败进度：统一使用全局进度
                if total_target_count:
                    print(f"✗ 第 {task[0]} 个网页提取失败: {e} ({task[0]}/{total_target_count})")
                else:
                    print(f"✗ 第 {task[0]} 个网页提取失败: {e} ({task[0]}/{total_count})")
                # 即使提取失败，也添加一个空结果以保持索引一致
                results.append((task[1], "", "", ""))
                
                # 实时更新全局变量，支持中断保存
                current_extracted_results = results.copy()
            
            # 每完成10个显示一次进度
            if completed_count % 10 == 0:
                # 显示进度：统一使用全局进度
                if total_target_count:
                    print(f"📊 提取进度: {task[0]}/{total_target_count} ({task[0]/total_target_count*100:.1f}%)")
                else:
                    print(f"📊 提取进度: {task[0]}/{total_count} ({task[0]/total_count*100:.1f}%)")
    
    # 更新全局状态
    current_extracted_results = results
    
    # 按原始顺序排序结果
    results.sort(key=lambda x: next(i for i, task in enumerate(extraction_tasks) if task[1] == x[0]))
    
    # 结果已在提取过程中实时显示，无需重复显示
    
    return results

def translate_single_result(result_data):
    # 翻译单个搜索结果（用于并发执行）
    if len(result_data) == 6:  # 包含发布时间
        i, url, title, text, description, publish_date = result_data
    else:  # 向后兼容
        i, url, title, text, description = result_data
        publish_date = ""
    
    print(f"\n--- 翻译第 {i} 个结果 ---")
    print(f"URL: {url}")
    
    # 翻译标题
    translated_title = translate_text(title, "title")
    
    # 翻译摘要
    translated_description = translate_text(description, "summary")
    
    # 先截取正文到500字符，然后翻译截取后的文本
    if CONTENT_EXTRACTION_CONFIG['enable_truncate']:
        truncated_text = truncate_text_smart(text, CONTENT_EXTRACTION_CONFIG['max_chars'])
        print(f"  正文已截取到 {len(truncated_text)} 字符")
        # 翻译截取后的文本
        translated_text = translate_text(truncated_text, "content")
    else:
        # 如果没有启用截取，仍然截取到500字符进行翻译
        truncated_text = truncate_text_smart(text, 500)
        print(f"  正文已截取到 {len(truncated_text)} 字符进行翻译")
        translated_text = translate_text(truncated_text, "content")
    
    print(f"第 {i} 个结果翻译完成")
    return (url, title, text, description, translated_title, translated_text, translated_description, publish_date)

def batch_translate_results(results):
    # 批量翻译搜索结果（并发翻译）
    global interrupted
    
    if not TRANSLATION_CONFIG['enable_translation']:
        print("翻译功能已禁用，跳过翻译")
        return results
    
    if interrupted:
        print("检测到中断信号，跳过翻译")
        return results
    
    print(f"\n开始并发翻译 {len(results)} 个搜索结果...")
    print(f"并发线程数: {TRANSLATION_CONFIG['max_workers']}")
    print(f"请求间隔: {TRANSLATION_CONFIG['request_delay']}秒")
    
    # 准备并发翻译的数据
    translation_tasks = []
    for i, result in enumerate(results):
        if len(result) == 5:  # 包含发布时间
            url, title, text, description, publish_date = result
            translation_tasks.append((i+1, url, title, text, description, publish_date))
        else:  # 向后兼容
            url, title, text, description = result
            translation_tasks.append((i+1, url, title, text, description))
    
    translated_results = []
    
    # 使用线程池进行并发翻译
    with concurrent.futures.ThreadPoolExecutor(max_workers=TRANSLATION_CONFIG['max_workers']) as executor:
        # 提交所有翻译任务
        future_to_result = {executor.submit(translate_single_result, task): task for task in translation_tasks}
        
        # 收集翻译结果
        for future in concurrent.futures.as_completed(future_to_result):
            # 检查是否被中断
            if interrupted:
                print("\n⚠ 检测到中断信号，停止翻译...")
                break
                
            try:
                result = future.result()
                translated_results.append(result)
            except Exception as e:
                print(f"翻译任务失败: {e}")
                # 如果翻译失败，添加原始数据
                original_task = future_to_result[future]
                if len(original_task) == 6:  # 包含发布时间
                    translated_results.append((original_task[1], original_task[2], original_task[3], original_task[4],
                                             original_task[2], original_task[3], original_task[4], original_task[5]))
                else:  # 向后兼容
                    translated_results.append((original_task[1], original_task[2], original_task[3], original_task[4],
                                             original_task[2], original_task[3], original_task[4], ""))
    
    # 按原始顺序排序结果
    translated_results.sort(key=lambda x: next(i for i, task in enumerate(translation_tasks) if task[1] == x[0]))
    
    print(f"\n并发翻译完成！共翻译了 {len(translated_results)} 个结果")
    return translated_results

def get_google_custom_search_results(query, num_results=None, total_target_count=None, start_index=0):
    global current_search_results, interrupted
    # 使用 Google官方API 获取搜索结果
    if not GOOGLE_API_KEY or not GOOGLE_SEARCH_ENGINE_ID:
        print("Google官方API 配置不完整，请手动填写search_config.xlsx中的 api_key 和 search_engine_id")
        return []
    
    # 使用配置中的默认值
    if num_results is None:
        num_results = SEARCH_CONFIG['num_results']
    
    # 检查Google API限制
    if num_results > 100:
        print(f"警告: Google官方API 每个关键词最多只能返回100个结果")
        print(f"您请求了 {num_results} 个结果，将限制为100个")
        num_results = 100
    
    results = []
    start_index = 1
    results_per_page = 10  # Google API 最大限制
    
    # 处理搜索语言 - 使用统一的语言处理函数
    search_lang, display_name = get_search_language_code('google')
    print(f"搜索语言: {display_name}")
    # 打印搜索国家
    country_display = get_country_display_name_from_excel(SEARCH_CONFIG['country'])
    print(f"搜索国家: Google默认" if not SEARCH_CONFIG['country'] or SEARCH_CONFIG['country'] == 'Google默认' else f"搜索国家: {country_display}")
    print(f"用户设置搜索结果数量: {num_results}")
    
    try:
        while len(results) < num_results and not interrupted:
            # 构建请求参数
            params = {
                'key': GOOGLE_API_KEY,
                'cx': GOOGLE_SEARCH_ENGINE_ID,
                'q': query,
                'start': start_index,
                'num': min(results_per_page, num_results - len(results)),
                'c2coff': '1'  # 关闭简体中文和繁体中文之间的自动转换
            }
            
            # 添加可选参数
            if search_lang:
                params['lr'] = search_lang
            if SEARCH_CONFIG['country'] and SEARCH_CONFIG['country'] != 'Google默认':
                params['cr'] = f"country{SEARCH_CONFIG['country']}"
            if EXCLUDE_TERMS:
                params['excludeTerms'] = EXCLUDE_TERMS
            
            # 发送请求，重试3次
            api_url = "https://www.googleapis.com/customsearch/v1"
            for attempt in range(3):
                try:
                    response = requests.get(api_url, params=params, timeout=30)
                    # 一行打印返回JSON中的 error.message（如存在），然后再抛出异常
                    print((((response.json() if 'application/json' in response.headers.get('Content-Type','').lower() else {}) or {}).get('error') or {}).get('message')) if response.status_code >= 400 else None
                    response.raise_for_status()
                    data = response.json()
                    break
                except Exception as e:
                    if attempt < 2:
                        print(f"Google API 请求失败 (尝试 {attempt + 1}/3): {e}")
                        time.sleep(2)
                    else:
                        print(f"Google API 请求最终失败: {e}")
                        raise e
            
            # 检查是否有错误
            if 'error' in data:
                print(f"API 错误: {data['error']['message']}")
                break
            
            # 检查是否有搜索结果
            if 'items' not in data or not data['items']:
                print("没有更多搜索结果")
                break
            
            # 处理搜索结果
            current_page_results = []
            for item in data['items']:
                if len(results) >= num_results:
                    break
                
                url = item.get('link', '')
                title = item.get('title', '')
                description = item.get('snippet', '')
                
                current_page_results.append((url, title, description))
                results.append((url, title, description))
                
                # 实时更新全局变量，支持中断保存
                current_search_results = results.copy()
            
            # 立即提取当前页的网页内容
            if current_page_results:
                current_page = (start_index - 1) // results_per_page + 1
                print(f"\n--- 开始提取第 {current_page} 页的网页内容 ---")
                extracted_results = extract_content_phase(current_page_results, total_target_count=total_target_count, start_index=start_index)
                # 将提取的内容合并到最终结果中
                for i, extracted_result in enumerate(extracted_results):
                    # 找到对应的原始结果并更新
                    for j, original_result in enumerate(results):
                        if original_result[0] == extracted_result[0]:  # URL匹配
                            results[j] = extracted_result
                            break
            
            # 检查是否还有更多结果
            search_info = data.get('searchInformation', {})
            total_results = int(search_info.get('totalResults', 0))
            current_start = int(search_info.get('startIndex', 1))
            
            # 使用实际获取到的结果数量
            current_page_results = len(data['items'])
            
            current_page = (start_index - 1) // results_per_page + 1
            
            # 只在第一页先打印总搜索结果数量
            if current_page == 1:
                print(f"总共约 {format_number(total_results)} 个搜索结果")
                print(f"第 {current_page} 页：已获取 {len(results)} 个结果")
            else:
                print(f"第 {current_page} 页：已获取 {len(results)} 个结果")
            
            # 如果当前页的结果数为0，说明没有更多结果了
            if current_page_results == 0:
                print("已获取所有可用结果")
                break
            
            # 如果当前页的结果数少于10个，但已经达到目标数量，则停止
            if current_page_results < results_per_page and len(results) >= num_results:
                print("已达到目标结果数量")
                break
            
            # 准备下一页
            start_index += results_per_page
            
            # 添加请求延迟，避免API限流（使用固定间隔，带默认值兜底）
            time.sleep(OTHER_CONFIG.get('sleep_interval', 1))
    
    except requests.exceptions.RequestException as e:
        print(f"API 请求失败: {e}")
    except Exception as e:
        print(f"获取搜索结果失败: {e}")
    
    print(f"搜索完成，共获取 {len(results)} 个结果")
    return results

def get_googlesearch_python_results(query, num_results=None, sleep_interval=None, total_target_count=None, start_index=0):
    """
    获取 Google 搜索结果前 num_results 个结果（包含URL、标题和摘要）
    使用googlesearch-python库，支持随机请求间隔控制（配置值±2秒范围内随机）
    """
    # 使用配置中的默认值
    if num_results is None:
        num_results = SEARCH_CONFIG['num_results']
    
    # 处理搜索语言 - 使用配置而不是自动检测
    search_lang, display_name = get_search_language_code('google')
    # 简化的语言参数处理 - 所有参数都始终存在，None表示Google默认设置
    if search_lang.startswith('lang_'):
        hl_param = search_lang.replace('lang_', '')  # lang_zh-CN -> zh-CN
        lr_param = search_lang  # lr使用原始格式
    elif search_lang == '' or search_lang == 'Google默认':
        hl_param = None  # Google默认设置
        lr_param = None    # Google默认设置
    else:
        hl_param = search_lang  # 保持原格式
        lr_param = None           # 其他情况lr为默认
    
    print(f"搜索语言: {display_name}")
    
    # 简化的国家参数处理 - 所有参数都始终存在，None表示Google默认设置
    if SEARCH_CONFIG['country'] and SEARCH_CONFIG['country'] != 'Google默认':
        gl_param = SEARCH_CONFIG['country'].lower()  # 小写国家代码
        cr_param = f"country{SEARCH_CONFIG['country']}"  # countryXX格式
    else:
        gl_param = None  # Google默认设置
        cr_param = None      # Google默认设置
    
    # 如果没有指定sleep_interval，使用其他配置中的固定值（带默认值兜底）
    if sleep_interval is None:
        sleep_interval = OTHER_CONFIG.get('sleep_interval', 1)
        print(f"使用配置的请求间隔: {sleep_interval}秒 (将在search函数内部随机化)")
    
    # 处理排除关键词 - 将排除关键词转换为Google搜索语法
    search_query = query
    if EXCLUDE_TERMS:
        exclude_list = EXCLUDE_TERMS.split()
        exclude_terms = " ".join([f"-{term}" for term in exclude_list])
        search_query = f"{query} {exclude_terms}"
        print(f"应用排除关键词: {exclude_terms}")
    
    results = []
    try:
        # 打印搜索国家
        country_display = get_country_display_name_from_excel(SEARCH_CONFIG['country'])
        print(f"搜索国家: Google默认" if not SEARCH_CONFIG['country'] or SEARCH_CONFIG['country'] == 'Google默认' else f"搜索国家: {country_display}")
        print(f"用户设置搜索结果数量: {num_results}")
        
        # 简化的搜索参数构建 - 所有参数都始终存在，None表示Google默认设置
        search_kwargs = {
            'q': search_query,  # 使用处理后的搜索查询（包含排除关键词）
            'num_results': num_results,
            'sleep_interval': sleep_interval,
            'advanced': True,
            'hl': hl_param,
            'gl': gl_param,
            'cr': cr_param,
            'lr': lr_param,
        }
        
        # 分批处理搜索结果，每10个结果提取一次内容
        batch_size = 10
        current_batch = []
        batch_count = 0
        
        for result in search(**search_kwargs):
            # 获取摘要信息，如果没有则使用空字符串
            description = getattr(result, 'description', '') or ''
            current_batch.append((result.url, result.title, description))
            results.append((result.url, result.title, description))
            
            # 当达到批次大小或达到目标数量时，立即提取内容
            if len(current_batch) >= batch_size or len(results) >= num_results:
                batch_count += 1
                print(f"\n--- 开始提取第 {batch_count} 批的网页内容 ---")
                extracted_results = extract_content_phase(current_batch, total_target_count=total_target_count, start_index=start_index)
                
                # 将提取的内容合并到最终结果中
                for extracted_result in extracted_results:
                    replaced = False
                    for j, original_result in enumerate(results):
                        if original_result[0] == extracted_result[0]:  # URL匹配
                            results[j] = extracted_result
                            replaced = True
                            # 继续查找并替换所有重复的URL
                            for k in range(j+1, len(results)):
                                if results[k][0] == extracted_result[0]:
                                    results[k] = extracted_result
                    if not replaced:
                        print(f"警告：未找到匹配的URL {extracted_result[0]}")
                
                # 清空当前批次
                current_batch = []
                
                # 如果已达到目标数量，停止搜索
                if len(results) >= num_results:
                    break
        
        # 处理最后一批不足10个的结果
        if current_batch:
            batch_count += 1
            print(f"\n--- 开始提取第 {batch_count} 批的网页内容 ---")
            extracted_results = extract_content_phase(current_batch, total_target_count=total_target_count, start_index=start_index)
            
            # 将提取的内容合并到最终结果中
            for extracted_result in extracted_results:
                replaced = False
                for j, original_result in enumerate(results):
                    if original_result[0] == extracted_result[0]:  # URL匹配
                        results[j] = extracted_result
                        replaced = True
                        # 继续查找并替换所有重复的URL
                        for k in range(j+1, len(results)):
                            if results[k][0] == extracted_result[0]:
                                results[k] = extracted_result
                if not replaced:
                    print(f"警告：未找到匹配的URL {extracted_result[0]}")
                        
    except Exception as e:
        print(f"获取搜索结果失败: {e}")
    
    return results

def get_serp_api_results(query, num_results=None, total_target_count=None, start_index=0):
    global current_search_results, interrupted
    # 使用 Serp API 获取搜索结果
    if not SERPAPI_AVAILABLE:
        print("Serp API 库未安装，请运行: pip install google-search-results")
        return []
    
    if not SERP_API_KEY:
        print("Serp API 密钥未配置，请在配置中填写 SERP_API_KEY")
        return []
    
    # 使用配置中的默认值
    if num_results is None:
        num_results = SEARCH_CONFIG['num_results']
    
    results = []
    start_index = 0
    results_per_page = 10  # Serp API 每页固定返回10个结果
    
    # 处理搜索查询 - 将排除关键词合并到查询中
    search_query = query
    if EXCLUDE_TERMS:
        # 将排除关键词转换为Google搜索语法：关键词 -排除词1 -排除词2
        exclude_list = EXCLUDE_TERMS.split()
        exclude_terms = " ".join([f"-{term}" for term in exclude_list])
        search_query = f"{query} {exclude_terms}"
    
    print(f"使用 Serp API 搜索: {search_query}")
    
    # 处理搜索语言 - 使用统一的语言处理函数
    search_lang, display_name = get_search_language_code('google')
    print(f"搜索语言: {display_name}")
    # 打印搜索国家
    country_display = get_country_display_name_from_excel(SEARCH_CONFIG['country'])
    print(f"搜索国家: Google默认" if not SEARCH_CONFIG['country'] or SEARCH_CONFIG['country'] == 'Google默认' else f"搜索国家: {country_display}")
    print(f"用户设置搜索结果数量: {num_results}")
    
    try:
        while len(results) < num_results and not interrupted:
            # 构建请求参数
            params = {
                "api_key": SERP_API_KEY,
                "engine": "google",
                "q": search_query,
                "google_domain": "google.com",
                "start": str(start_index)
            }
            
            # 添加语言参数
            if search_lang:
                params["lr"] = search_lang
            
            # 添加国家参数
            if SEARCH_CONFIG['country'] and SEARCH_CONFIG['country'] != 'Google默认':
                params["cr"] = f"country{SEARCH_CONFIG['country']}"
            
            # 发送请求，重试3次
            for attempt in range(3):
                try:
                    search = GoogleSearch(params)
                    search_results = search.get_dict()
                    break
                except Exception as e:
                    if attempt < 2:
                        print(f"Serp API 请求失败 (尝试 {attempt + 1}/3): {e}")
                        time.sleep(2)
                    else:
                        print(f"Serp API 请求最终失败: {e}")
                        raise e
            
            # 检查是否有错误
            if 'error' in search_results:
                print(f"Serp API 错误: {search_results['error']}")
                break
            
            # 检查是否有搜索结果
            if 'organic_results' not in search_results or not search_results['organic_results']:
                print("没有更多搜索结果")
                break
            
            # 处理搜索结果
            current_page_results = []
            for item in search_results['organic_results']:
                if len(results) >= num_results:
                    break
                
                url = item.get('link', '')
                title = item.get('title', '')
                description = item.get('snippet', '')
                
                current_page_results.append((url, title, description))
                results.append((url, title, description))
                
                # 实时更新全局变量，支持中断保存
                current_search_results = results.copy()
            
            # 立即提取当前页的网页内容
            if current_page_results:
                current_page = start_index // results_per_page + 1
                print(f"\n--- 开始提取第 {current_page} 页的网页内容 ---")
                extracted_results = extract_content_phase(current_page_results, total_target_count=total_target_count, start_index=start_index)
                # 将提取的内容合并到最终结果中
                for i, extracted_result in enumerate(extracted_results):
                    # 找到对应的原始结果并更新
                    for j, original_result in enumerate(results):
                        if original_result[0] == extracted_result[0]:  # URL匹配
                            results[j] = extracted_result
                            break
            
            # 检查是否还有更多结果
            search_info = search_results.get('search_information', {})
            total_results = search_info.get('total_results', 0)
            
            current_page = start_index // results_per_page + 1
            current_page_results = len(search_results['organic_results'])
            
            # 只在第一页先打印总搜索结果数量
            if current_page == 1:
                print(f"总共约 {format_number(total_results)} 个搜索结果")
                print(f"第 {current_page} 页：已获取 {len(results)} 个结果")
            else:
                print(f"第 {current_page} 页：已获取 {len(results)} 个结果")
            
            # 如果当前页的结果数为0，说明没有更多结果了
            if current_page_results == 0:
                print("已获取所有可用结果")
                break
            
            # 如果当前页的结果数少于10个，但已经达到目标数量，则停止
            if current_page_results < results_per_page and len(results) >= num_results:
                print("已达到目标结果数量")
                break
            
            # 准备下一页
            start_index += results_per_page
            
            # 添加请求延迟，避免API限流（带默认值兜底）
            time.sleep(OTHER_CONFIG.get('sleep_interval', 1))
    
    except Exception as e:
        print(f"Serp API 请求失败: {e}")
    
    print(f"搜索完成，共获取 {len(results)} 个结果")
    return results

def get_search_results(query, num_results=None, total_target_count=None, start_index=0):
    # 根据配置选择搜索方法
    if SEARCH_METHOD == 'google_api':
        return get_google_custom_search_results(query, num_results, total_target_count, start_index)
    elif SEARCH_METHOD == 'serp_api':
        return get_serp_api_results(query, num_results, total_target_count, start_index)
    else:
        return get_googlesearch_python_results(query, num_results, total_target_count=total_target_count, start_index=start_index)

def extract_with_newspaper(url):
    # 使用newspaper3k提取正文，通过requests控制超时
    try:
        # 使用requests先下载网页内容，控制超时时间
        network_timeout = CONTENT_EXTRACTION_CONFIG.get('network_timeout', 5)
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # 使用配置中的网络超时时间
        response = requests.get(url, headers=headers, timeout=network_timeout)
        response.raise_for_status()
        
        # 使用newspaper3k解析已下载的内容
        article = Article(url)
        article.set_html(response.text)
        article.parse()
        return article.text
    except Exception as e:
        print(f"newspaper3k提取失败: {url}, 错误: {e}")
        return ""

def extract_with_trafilatura(html_content, url=""):
    # 使用trafilatura提取正文和元数据（包括发布时间）
    if not TRAFILATURA_AVAILABLE:
        return "", ""
    
    try:
        if html_content:
            # 从已下载的HTML内容中提取正文
            content = trafilatura.extract(html_content, url=url)
            
            # 提取元数据，包括发布时间
            metadata = trafilatura.extract_metadata(html_content)
            publish_date = ""
            if metadata and hasattr(metadata, 'date'):
                publish_date = metadata.date
            
            return content or "", publish_date or ""
        else:
            return "", ""
    except Exception as e:
        print(f"trafilatura提取失败: {url}, 错误: {e}")
        return "", ""



def extract_article_content_complete(url, page_idx: int = None, total_pages: int = None):
    """
    统一提取网页标题、描述、正文和发布时间 - 优化版（一次下载，多种解析方法）
    返回: (title, content, description, publish_date)
    """
    import time
    from urllib.parse import urlparse
    
    # 日志上下文：序号/总数、域名
    domain = urlparse(url).netloc
    ctx = f"[{page_idx}/{total_pages}]" if page_idx and total_pages else ""
    
    # 第一步：下载网页内容（只请求一次网络）
    try:
        network_timeout = CONTENT_EXTRACTION_CONFIG.get('network_timeout', 5)
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        }
        
        start_time = time.time()
        response = requests.get(url, headers=headers, timeout=network_timeout)
        response.raise_for_status()
        
        html_content = response.text
        if not html_content:
            return "", "", "", ""
        
        # 智能编码检测：优先使用响应头编码，仅在必要时使用chardet检测
        html_content = smart_encoding_detection(html_content, response)
            
    except Exception as e:
        return "", "", "", ""
    
    # 第二步：从HTML中提取标题和描述
    title = ""
    description = ""
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 提取标题
        title_tag = soup.find('title')
        if title_tag:
            title = title_tag.get_text()
            title = html.unescape(title)
            import re
            title = re.sub(r'\s+', ' ', title).strip()
        
        # 提取description
        desc_selectors = [
            'meta[name="description"]',
            'meta[property="og:description"]',
            'meta[name="twitter:description"]'
        ]
        
        for selector in desc_selectors:
            desc_tag = soup.select_one(selector)
            if desc_tag and desc_tag.get('content'):
                description = desc_tag.get('content')
                description = html.unescape(description)
                description = re.sub(r'\s+', ' ', description).strip()
                break
                
    except Exception as e:
        pass
    
    # 第三步：提取发布时间
    publish_date = ""
    try:
        if TRAFILATURA_AVAILABLE:
            metadata = trafilatura.extract_metadata(html_content)
            if metadata and hasattr(metadata, 'date') and metadata.date:
                publish_date = metadata.date
    except Exception as e:
        pass
    
    # 第四步：用不同方法解析同一份HTML内容提取正文
    best_content = ""
    best_method = ""
    best_length = 0
    
    # 尝试trafilatura解析
    try:
        if TRAFILATURA_AVAILABLE:
            content = trafilatura.extract(html_content, url=url)
            
            if content and len(content) > 100:
                content_length = len(content)
                if content_length > best_length:
                    best_content = content
                    best_method = "trafilatura"
                    best_length = content_length
    except Exception as e:
        pass
    
    # 尝试newspaper3k解析
    try:
        article = Article(url)
        article.set_html(html_content)
        article.parse()
        content = article.text
        
        if content and len(content) > 100:
            content_length = len(content)
            if content_length > best_length:
                best_content = content
                best_method = "newspaper3k"
                best_length = content_length
    except Exception as e:
        pass
    
    # 返回标题、正文、描述、发布时间
    return title, best_content, description, publish_date

def extract_article_text_robust(url, page_idx: int = None, total_pages: int = None):
    # 多库组合提取正文 - 优化版（一次下载，多种解析方法）
    _, content, _, _ = extract_article_content_complete(url, page_idx, total_pages)
    return content

def extract_article_text(url, page_idx: int = None, total_pages: int = None):
    # 提取网页正文 - 保持向后兼容
    return extract_article_text_robust(url, page_idx=page_idx, total_pages=total_pages)

def truncate_text_smart(text, max_chars=500):
    # 智能截取文本，保持句子完整性
    if not text or len(text) <= max_chars:
        return text
    
    # 截取到指定长度
    truncated = text[:max_chars]
    
    # 查找最后一个完整的句子结束符
    sentence_endings = ['.', '!', '?', '。', '！', '？']
    last_sentence_end = -1
    
    for ending in sentence_endings:
        pos = truncated.rfind(ending)
        if pos > last_sentence_end:
            last_sentence_end = pos
    
    # 如果找到句子结束符，在它后面截断
    if last_sentence_end > max_chars * 0.7:  # 确保不会截掉太多内容
        return truncated[:last_sentence_end + 1]
    else:
        # 如果没找到合适的句子结束符，在单词边界截断
        last_space = truncated.rfind(' ')
        if last_space > max_chars * 0.8:
            return truncated[:last_space] + "..."
        else:
            return truncated + "..."

def extract_article_title_and_text(url, google_title="", google_description=""):
    """
    提取网页标题、description、正文和发布时间（使用多库组合），不进行翻译
    优化版：只进行一次网络请求
    """
    try:
        # 使用统一函数一次性提取标题、正文、描述和发布时间
        html_title, text, html_description, publish_date = extract_article_content_complete(url)
        
        # 确定最终使用的标题
        title = html_title if html_title else google_title
        if not title and google_title:
            title = google_title
            print(f"使用Google搜索结果标题作为备选: {title}")
        
        # 确定最终使用的description（优先级：HTML > Google）
        final_description = html_description if html_description else google_description
        
        return title, text, final_description, publish_date
    except Exception as e:
        print(f"提取失败: {url}, 错误: {e}")
        return "", "", "", ""

def save_results_to_pdf(results, filename="results.pdf"):
    """
    将搜索结果保存为 PDF 文件，支持中文
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # 添加支持中文的字体
    pdf.add_font('ArialUnicode', '', 'msyh.ttf', uni=True)
    pdf.set_font('ArialUnicode', size=12)

    for i, (url, title, text) in enumerate(results, 1):
        pdf.set_font('ArialUnicode', style='B', size=12)
        pdf.cell(0, 10, f"第 {i} 个搜索结果", ln=True)
        pdf.set_font('ArialUnicode', size=12)
        pdf.cell(0, 10, f"标题: {title}", ln=True)
        pdf.cell(0, 10, f"URL: {url}", ln=True)
        pdf.multi_cell(0, 10, f"正文内容: {text}")
        pdf.ln(10)

    pdf.output(filename)

def add_hyperlink(paragraph, text, url):
    # 在段落中添加超链接，并设置蓝色字体和下划线样式
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    # 设置蓝色字体和下划线样式
    color = OxmlElement("w:color")
    color.set("val", "0000FF")  # 蓝色
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set("val", "single")  # 下划线
    r_pr.append(underline)

    new_run.append(r_pr)

    new_text = OxmlElement("w:t")
    new_text.text = text
    new_run.append(new_text)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def extract_domain(url):
    # 从URL中提取域名
    try:
        parsed_url = urlparse(url)
        domain = parsed_url.netloc
        # 移除www前缀
        if domain.startswith('www.'):
            domain = domain[4:]
        return domain
    except Exception:
        return url

def save_results_to_docx(results, filename="results.docx", keywords_list=None, languages_list=None, countries_list=None, requested_count=None):
    """
    将搜索结果保存为 DOCX 文件，标题添加超链接，支持翻译内容
    """
    doc = Document()
    keywords_display, languages_display, countries_display = get_unified_search_info_display(keywords_list, languages_list, countries_list)
    doc.add_heading(f"搜索结果 - {keywords_display} ({__import__('datetime').datetime.now().strftime('%Y-%m-%d')})", level=1)
    doc.add_paragraph(f"搜索关键词: {keywords_display}")
    doc.add_paragraph(f"搜索语种: {languages_display}")
    doc.add_paragraph(f"搜索国家: {countries_display}")
    
    # 显示搜索结果数量信息
    actual_count = len(results)
    if requested_count is not None:
        doc.add_paragraph(f"搜索结果数量: 用户请求 {requested_count} 个，实际获取 {actual_count} 个")
    else:
        doc.add_paragraph(f"搜索结果数量: {actual_count} 个")

    for i, result in enumerate(results, 1):
        # 解包结果，支持翻译内容和发布时间
        if len(result) == 8:  # 包含翻译内容和发布时间
            url, title, text, description, translated_title, translated_text, translated_description, publish_date = result
        elif len(result) == 7:  # 包含翻译内容，但不包含发布时间（向后兼容）
            url, title, text, description, translated_title, translated_text, translated_description = result
            publish_date = ""
        elif len(result) == 5:  # 包含发布时间，但不包含翻译内容
            url, title, text, description, publish_date = result
            translated_title, translated_text, translated_description = title, text, description
        elif len(result) == 4:  # 不包含翻译内容和发布时间（向后兼容）
            url, title, text, description = result
            translated_title, translated_text, translated_description = title, text, description
            publish_date = ""
        elif len(result) == 3:  # 只有基本搜索结果（URL, title, description）
            url, title, description = result
            text = ""
            translated_title, translated_text, translated_description = title, text, description
            publish_date = ""
        else:
            print(f"警告: 结果格式不正确，包含 {len(result)} 个元素: {result}")
            continue
        
        doc.add_heading(f"第 {i} 个搜索结果", level=2)
        
        # 添加标题（原始标题带超链接，翻译标题紧跟其后）
        title_paragraph = doc.add_paragraph()
        add_hyperlink(title_paragraph, title, url)
        
        # 添加翻译后标题（如果启用翻译），使用相同样式
        if TRANSLATION_CONFIG['enable_translation'] and translated_title != title and translated_title:
            # 在同一个段落中添加翻译标题
            title_paragraph.add_run(f" {translated_title}")
        
        # 添加URL信息
        doc.add_paragraph(f"网页URL: {url}")
        
        # 添加公司网址信息
        domain = extract_domain(url)
        doc.add_paragraph(f"公司网址: {domain}")
        
        # 添加发布时间信息（如果有的话）
        if publish_date:
            doc.add_paragraph(f"发布时间: {publish_date}")
        
        # 添加网站摘要信息
        if description:
            doc.add_paragraph(f"摘要: {description}")
        
        # 添加翻译后摘要（如果启用翻译）
        if TRANSLATION_CONFIG['enable_translation'] and translated_description != description and translated_description:
            doc.add_paragraph(translated_description)
        
        # 添加正文内容（智能截取）
        doc.add_paragraph("正文内容:")
        if CONTENT_EXTRACTION_CONFIG['enable_truncate']:
            truncated_text = truncate_text_smart(text, CONTENT_EXTRACTION_CONFIG['max_chars'])
            doc.add_paragraph(truncated_text)
        else:
            doc.add_paragraph(text)
        
        # 添加翻译后正文内容（如果启用翻译）
        if TRANSLATION_CONFIG['enable_translation'] and translated_text != text and translated_text:
            if CONTENT_EXTRACTION_CONFIG['enable_truncate']:
                truncated_translated_text = truncate_text_smart(translated_text, CONTENT_EXTRACTION_CONFIG['max_chars'])
                doc.add_paragraph(truncated_translated_text)
            else:
                doc.add_paragraph(translated_text)

    doc.save(filename)

def save_results_to_html(results, filename="results.html", keywords_list=None, languages_list=None, countries_list=None, requested_count=None):
    """
    将搜索结果保存为简洁的HTML文件
    """
    html_content = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>搜索结果 - {SEARCH_KEYWORD}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            background: #fff;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        
        h1 {{
            font-size: 2em;
            font-weight: 600;
            margin-bottom: 20px;
            color: #000;
            border-bottom: 2px solid #e1e4e8;
            padding-bottom: 10px;
        }}
        
        .search-info {{
            background: #f6f8fa;
            border: 1px solid #d0d7de;
            border-radius: 6px;
            padding: 16px;
            margin-bottom: 30px;
            font-size: 14px;
            color: #656d76;
        }}
        
        .result-item {{
            margin-bottom: 40px;
            padding-bottom: 30px;
            border-bottom: 1px solid #e1e4e8;
        }}
        
        .result-item:last-child {{
            border-bottom: none;
        }}
        
        .result-number {{
            display: inline-block;
            background: #0969da;
            color: white;
            width: 24px;
            height: 24px;
            border-radius: 12px;
            text-align: center;
            line-height: 24px;
            font-size: 12px;
            font-weight: 600;
            margin-right: 12px;
            vertical-align: top;
        }}
        
        .result-title {{
            font-size: 1.25em;
            font-weight: 600;
            margin-bottom: 8px;
            line-height: 1.3;
        }}
        
        .result-title a {{
            color: #0969da;
            text-decoration: none;
        }}
        
        .result-title a:hover {{
            text-decoration: underline;
        }}
        
        .result-title-link {{
            color: #0969da;
            text-decoration: none;
        }}
        
        .result-title-link:hover {{
            text-decoration: underline;
        }}
        
        .result-url {{
            color: #656d76;
            font-size: 14px;
            margin-bottom: 8px;
            word-break: break-all;
        }}
        
        .result-url a {{
            color: #0969da;
            text-decoration: none;
        }}
        
        .result-url a:hover {{
            text-decoration: underline;
        }}
        
        .domain-badge {{
            display: inline-block;
            background: #dbeafe;
            color: #1e40af;
            padding: 2px 6px;
            border-radius: 3px;
            font-size: 12px;
            font-weight: 500;
            margin-left: 8px;
        }}
        
        .publish-date {{
            display: inline-block;
            background: #f3f4f6;
            color: #6b7280;
            padding: 2px 6px;
            border-radius: 3px;
            font-size: 12px;
            font-weight: 500;
            margin-left: 8px;
        }}
        
        .result-content {{
            margin-top: 12px;
            font-size: 14px;
            line-height: 1.5;
            color: #656d76;
        }}
        
        .result-content h4 {{
            color: #24292f;
            font-size: 14px;
            font-weight: 600;
            margin: 16px 0 8px 0;
        }}
        
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #e1e4e8;
            font-size: 12px;
            color: #656d76;
            text-align: center;
        }}
        
        @media (max-width: 768px) {{
            body {{
                padding: 16px;
            }}
            
            h1 {{
                font-size: 1.5em;
            }}
            
            .result-title {{
                font-size: 1.1em;
            }}
        }}
    </style>
</head>
<body>
    <h1>搜索结果-{get_unified_search_info_display(keywords_list, languages_list, countries_list)[0]} ({__import__('datetime').datetime.now().strftime('%Y-%m-%d')})</h1>
    
    <div class="search-info">
        <strong>搜索关键词:</strong> {get_unified_search_info_display(keywords_list, languages_list, countries_list)[0]} | 
        <strong>搜索语种:</strong> {get_unified_search_info_display(keywords_list, languages_list, countries_list)[1]} | 
        <strong>搜索国家:</strong> {get_unified_search_info_display(keywords_list, languages_list, countries_list)[2]} | 
        <strong>搜索结果数量:</strong> {f"用户请求 {requested_count} 个，实际获取 {len(results)} 个" if requested_count is not None else f"{len(results)} 个"}
    </div>
"""

    for i, result in enumerate(results, 1):
        # 解包结果，支持翻译内容和发布时间
        if len(result) == 8:  # 包含翻译内容和发布时间
            url, title, text, description, translated_title, translated_text, translated_description, publish_date = result
        elif len(result) == 7:  # 包含翻译内容，但不包含发布时间（向后兼容）
            url, title, text, description, translated_title, translated_text, translated_description = result
            publish_date = ""
        elif len(result) == 5:  # 包含发布时间，但不包含翻译内容
            url, title, text, description, publish_date = result
            translated_title, translated_text, translated_description = title, text, description
        elif len(result) == 4:  # 不包含翻译内容和发布时间（向后兼容）
            url, title, text, description = result
            translated_title, translated_text, translated_description = title, text, description
            publish_date = ""
        elif len(result) == 3:  # 只有基本搜索结果（URL, title, description）
            url, title, description = result
            text = ""
            translated_title, translated_text, translated_description = title, text, description
            publish_date = ""
        else:
            print(f"警告: 结果格式不正确，包含 {len(result)} 个元素: {result}")
            continue
        
        domain = extract_domain(url)
        html_content += f"""
    <div class="result-item">
        <div class="result-title">
            <span class="result-number">{i}</span>
            <a href="{url}" target="_blank">{title}</a>
            {f' <a href="{url}" target="_blank" class="result-title-link">{translated_title}</a>' if TRANSLATION_CONFIG['enable_translation'] and translated_title != title and translated_title else ''}
        </div>
        <div class="result-url">
            <a href="{url}" target="_blank">{url}</a>
            <span class="domain-badge">{domain}</span>
            {f'<span class="publish-date">发布时间: {publish_date}</span>' if publish_date else ''}
        </div>
        {f'<div class="result-content"><p><strong>摘要:</strong> {description}</p></div>' if description else ''}
        {f'<div class="result-content"><p>{translated_description}</p></div>' if TRANSLATION_CONFIG['enable_translation'] and translated_description != description and translated_description else ''}
        <div class="result-content">
            <h4>正文内容:</h4>
            <p>{truncate_text_smart(text, CONTENT_EXTRACTION_CONFIG['max_chars']) if CONTENT_EXTRACTION_CONFIG['enable_truncate'] else text}</p>
        </div>
        {f'<div class="result-content"><p>{truncate_text_smart(translated_text, CONTENT_EXTRACTION_CONFIG["max_chars"]) if CONTENT_EXTRACTION_CONFIG["enable_truncate"] else translated_text}</p></div>' if TRANSLATION_CONFIG['enable_translation'] and translated_text != text and translated_text else ''}
    </div>
"""

    html_content += f"""
    <div class="footer">
        <p>搜索结果由 Google 搜索生成 | 生成时间: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
</body>
</html>
"""

    with open(filename, 'w', encoding='utf-8') as f:
        f.write(html_content)

def save_results_to_excel(results, filename="results.xlsx", search_keyword=None, search_keywords_per_row=None, search_country=None, search_countries_per_row=None, search_languages_per_row=None):
    """
    将搜索结果保存为 Excel 文件
    """
    # 使用传入的关键词/国家或全局配置
    if search_keyword is None and not search_keywords_per_row:
        search_keyword = SEARCH_KEYWORD
    if search_country is None and not search_countries_per_row:
        search_country = SEARCH_CONFIG.get('country', '')
    
    # 准备数据列表
    data = []
    
    for i, result in enumerate(results, 1):
        # 解包结果，支持翻译内容和发布时间
        if len(result) == 8:  # 包含翻译内容和发布时间
            url, title, text, description, translated_title, translated_text, translated_description, publish_date = result
        elif len(result) == 7:  # 包含翻译内容，但不包含发布时间（向后兼容）
            url, title, text, description, translated_title, translated_text, translated_description = result
            publish_date = ""
        elif len(result) == 5:  # 包含发布时间，但不包含翻译内容
            url, title, text, description, publish_date = result
            translated_title, translated_text, translated_description = title, text, description
        elif len(result) == 4:  # 不包含翻译内容和发布时间（向后兼容）
            url, title, text, description = result
            translated_title, translated_text, translated_description = title, text, description
            publish_date = ""
        elif len(result) == 3:  # 只有基本搜索结果（URL, title, description）
            url, title, description = result
            text = ""
            translated_title, translated_text, translated_description = title, text, description
            publish_date = ""
        else:
            print(f"警告: 结果格式不正确，包含 {len(result)} 个元素: {result}")
            continue
        
        # 提取域名
        domain = extract_domain(url)
        
        # 截取正文内容
        if CONTENT_EXTRACTION_CONFIG['enable_truncate']:
            truncated_text = truncate_text_smart(text, CONTENT_EXTRACTION_CONFIG['max_chars'])
            truncated_translated_text = truncate_text_smart(translated_text, CONTENT_EXTRACTION_CONFIG['max_chars']) if TRANSLATION_CONFIG['enable_translation'] else ""
        else:
            truncated_text = text
            truncated_translated_text = translated_text if TRANSLATION_CONFIG['enable_translation'] else ""
        
        # 选择该行的搜索关键词
        if search_keywords_per_row and len(search_keywords_per_row) >= i:
            row_keyword = search_keywords_per_row[i-1]
        else:
            row_keyword = search_keyword
        
        # 选择该行的搜索国家
        if search_countries_per_row and len(search_countries_per_row) >= i:
            row_country = search_countries_per_row[i-1]
        else:
            row_country = search_country
        
        # 处理空字符串的搜索国家，显示为"Google默认"
        if not row_country:
            row_country = "Google默认"
        
        # 从搜索国家代码工作表获取国家显示名称
        row_country_display = get_country_display_name_from_excel(row_country)
        
        # 选择该行的搜索语言
        if search_languages_per_row and len(search_languages_per_row) >= i:
            row_language = search_languages_per_row[i-1]
            row_language_display = get_language_display_name_from_excel(row_language)
        else:
            # 使用全局配置获取搜索语言显示名称
            search_lang, search_lang_display = get_search_language_code('google')
            if not search_lang:
                row_language_display = get_language_display_name_from_excel('Google默认')
            else:
                row_language_display = get_language_display_name_from_excel(search_lang)
        
        # 添加数据行（第二列为搜索关键词，第三列为搜索国家，第四列为搜索语言）
        row_data = {
            '序号': i,
            '搜索关键词': row_keyword,
            '搜索国家': row_country_display,
            '搜索语言': row_language_display,
            '原始标题': title,
            '翻译标题': translated_title if TRANSLATION_CONFIG['enable_translation'] and translated_title != title else "",
            '网页URL': url,
            '公司网址': domain,
            '发布时间': publish_date,
            '原始摘要': description,
            '翻译摘要': translated_description if TRANSLATION_CONFIG['enable_translation'] and translated_description != description else "",
            '原始正文': truncated_text,
            '翻译正文': truncated_translated_text if TRANSLATION_CONFIG['enable_translation'] and translated_text != text else ""
        }
        data.append(row_data)
    
    # 创建DataFrame
    df = pd.DataFrame(data)
    
    # 保存为Excel文件
    with pd.ExcelWriter(filename, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='搜索结果', index=False)
        
        # 获取工作表对象以调整列宽和字体
        worksheet = writer.sheets['搜索结果']
        
        # 导入字体样式
        from openpyxl.styles import Font
        
        # 设置等线字体样式
        default_font = Font(name='等线', size=11)
        header_font = Font(name='等线', size=11, bold=True)
        
        # 设置所有单元格的字体为等线
        for row in worksheet.iter_rows():
            for cell in row:
                if cell.row == 1:  # 标题行使用粗体
                    cell.font = header_font
                else:  # 数据行使用普通字体
                    cell.font = default_font
        
        # 自动调整列宽，并为特定列设置宽度
        column_min_widths = {
            '搜索语言': 11,  # 搜索语言列最小宽度11
            '搜索国家': 11,  # 搜索国家列最小宽度11
            '翻译标题': 30,  # 翻译标题列最小宽度30
            '翻译摘要': 35,  # 翻译摘要列最小宽度35
            '翻译正文': 40,  # 翻译正文列最小宽度40
        }
        
        # 固定宽度的列（不根据内容调整）
        column_fixed_widths = {
            '搜索关键词': 40,  # 搜索关键词列固定宽度40
            '公司网址': 30,    # 公司网址列固定宽度30
            '网页URL': 30,     # 网页URL列固定宽度30
        }
        
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            column_name = column[0].value  # 获取列名（标题行的值）
            
            # 如果是固定宽度的列，直接设置固定宽度
            if column_name in column_fixed_widths:
                adjusted_width = column_fixed_widths[column_name]
            else:
                # 对于其他列，计算基于内容的宽度
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                
                # 计算调整后的宽度
                adjusted_width = min(max_length + 2, 50)  # 限制最大宽度为50
                
                # 如果是有最小宽度要求的列，应用最小宽度限制
                if column_name in column_min_widths:
                    adjusted_width = max(adjusted_width, column_min_widths[column_name])
            
            worksheet.column_dimensions[column_letter].width = adjusted_width

def load_multi_keywords_tasks(filename: str = "search_config.xlsx"):
    """
    从"多搜索关键词模式"工作表加载任务列表。
    返回列表，每项为字典：
      { 'keyword': str, 'exclude': str, 'num_results': int|None, 'language': str, 'country': str }
    空行或无"搜索关键词"的行会被跳过。
    """
    tasks = []
    try:
        df = pd.read_excel(filename, sheet_name='多搜索关键词模式')
    except Exception:
        return tasks

    # 兼容列名（用户可能手动改动）
    rename_map = {}
    if '搜索关键词' not in df.columns:
        for col in df.columns:
            if str(col).strip() in ['关键词', '关键字', 'keyword']:
                rename_map[col] = '搜索关键词'
    if rename_map:
        df = df.rename(columns=rename_map)

    for _, row in df.iterrows():
        kw = normalize_cell(row.get('搜索关键词', ''))
        if not kw:
            continue
        exclude = normalize_cell(row.get('排除关键词', '')) if '排除关键词' in df.columns else ''
        # 数量
        nr = parse_int_or_none(row.get('搜索结果数量', None))
        language = normalize_cell(row.get('搜索语言', '')) if '搜索语言' in df.columns else ''
        country = normalize_cell(row.get('搜索国家', '')) if '搜索国家' in df.columns else ''

        tasks.append({
            'keyword': kw,
            'exclude': exclude,
            'num_results': nr,
            'language': language,
            'country': country
        })
    return tasks

def process_one_query(query: str, total_target_count=None, start_index=0):
    """
    处理单个查询的完整流程：搜索→并发提取→批量翻译。
    返回 translated_results 列表。
    
    Args:
        query: 搜索关键词
        total_target_count: 多关键词模式下的总目标数量（仅用于进度显示）
        start_index: 多关键词模式下的起始索引（用于进度显示）
    """
    global current_search_results, current_extracted_results, current_translated_results, interrupted

    # 使用新的边搜索边提取逻辑，每个查询使用自己的num_results
    search_results = get_search_results(query, num_results=SEARCH_CONFIG['num_results'], total_target_count=total_target_count, start_index=start_index)
    current_search_results = search_results
    if not search_results:
        return []

    # 搜索结果已经包含了提取的内容，直接使用
    results = search_results
    current_extracted_results = results

    if interrupted:
        return results

    translated_results = batch_translate_results(results)
    current_translated_results = translated_results
    return translated_results

def send_email_with_attachment(html_filename, docx_filename, xlsx_filename, subject="搜索结果报告", num_results=0):
    """
    通过QQ邮箱发送HTML内容作为邮件正文，并附加Word文档和Excel文件
    """
    if not EMAIL_CONFIG['enable_email']:
        print("邮件发送功能已禁用，请在配置中启用。")
        return False
    
    try:
        # 读取HTML文件内容
        with open(html_filename, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # 创建邮件对象
        msg = MIMEMultipart('mixed')
        msg['From'] = EMAIL_CONFIG['sender_email']
        msg['To'] = EMAIL_CONFIG['recipient_email']
        msg['Subject'] = f"{subject} - {SEARCH_KEYWORD} ({__import__('datetime').datetime.now().strftime('%Y-%m-%d')})"
        
        # 纯文本版本
        text_body = f"""
        您好！
        
        这是关于"{SEARCH_KEYWORD}"的搜索结果报告。
        
        报告包含：
        - {num_results}个相关网站的详细信息
        - 网站标题、URL、域名和摘要
        - 完整的网页正文内容
        - 中文翻译版本
        
        生成时间：{__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        
        此邮件由自动搜索脚本生成。
        
        附件说明：
        - Word文档：包含完整的搜索结果和翻译，适合阅读和打印
        - Excel文件：包含结构化的数据表格，适合数据分析和处理
        
        注意：此邮件包含HTML格式的搜索结果和多个附件，请确保您的邮件客户端支持HTML显示和附件下载。
        """
        
        # 创建纯文本和HTML部分
        text_part = MIMEText(text_body, 'plain', 'utf-8')
        html_part = MIMEText(html_content, 'html', 'utf-8')
        
        # 创建多部分容器
        msg_alternative = MIMEMultipart('alternative')
        msg_alternative.attach(text_part)
        msg_alternative.attach(html_part)
        msg.attach(msg_alternative)
        
        # 添加Word文档附件
        try:
            with open(docx_filename, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename= {docx_filename}'
                )
                msg.attach(part)
            print(f"Word文档附件已添加: {docx_filename}")
        except FileNotFoundError:
            print(f"警告: 找不到Word文档文件 {docx_filename}")
        except Exception as e:
            print(f"添加Word文档附件失败: {e}")
        
        # 添加Excel文件附件
        try:
            with open(xlsx_filename, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename= {xlsx_filename}'
                )
                msg.attach(part)
            print(f"Excel文件附件已添加: {xlsx_filename}")
        except FileNotFoundError:
            print(f"警告: 找不到Excel文件 {xlsx_filename}")
        except Exception as e:
            print(f"添加Excel文件附件失败: {e}")
        
        # 连接SMTP服务器并发送邮件
        server = smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port'])
        server.starttls()  # 启用TLS加密
        server.login(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['sender_auth_code'])
        
        text = msg.as_string()
        server.sendmail(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['recipient_email'], text)
        server.quit()
        
        print(f"邮件发送成功！收件人：{EMAIL_CONFIG['recipient_email']}")
        print("HTML内容已作为邮件正文发送，Word文档和Excel文件已作为附件发送。")
        return True
        
    except Exception as e:
        print(f"邮件发送失败：{e}")
        return False

def main():
    # 注册信号处理器
    signal.signal(signal.SIGINT, signal_handler)
    
    # 注册程序退出时的紧急保存处理器
    atexit.register(emergency_exit_handler)
    
    # 检查配置文件
    print("="*60)
    print("检查配置文件...")
    print("="*60)
    
    config_exists = check_and_create_config_file()
    
    if config_exists:
        # 如果配置文件存在，尝试从Excel文件加载配置
        if load_config_from_excel():
            print("✓ 配置加载成功，将使用Excel文件中的配置")
        else:
            print("⚠ 配置加载失败，将使用代码中的默认配置")
    else:
        print("\n📋 请按照以下步骤配置：")
        print("1. 打开 search_config.xlsx 文件")
        print("2. 在'配置模板'工作表中填写搜索关键词等配置")
        print("3. 推荐选择'googlesearch'模式（无需API）")
        print("4. 保存文件后重新运行程序")
        print("\n" + "="*50)
        input("配置完成后，按回车键退出程序...")
        return False  # 返回False表示不需要最后的暂停
    
    print("="*60)
    print("开始执行搜索任务...")
    print("="*60)
    
    # 如果硬编码开关启用，则走多关键词批量模式
    if ENABLE_MULTI_KEYWORDS_MODE:
        print("启用多搜索关键词模式：从工作表读取批量任务")
        tasks = load_multi_keywords_tasks("search_config.xlsx")
        if not tasks:
            print("未在'多搜索关键词模式'工作表中找到任务，回退到单关键词模式")
        else:
            # 计算所有关键词的总目标数量
            total_target_count = sum(t['num_results'] if t['num_results'] is not None else SEARCH_CONFIG['num_results'] for t in tasks)
            print(f"多关键词搜索总目标数量: {total_target_count}")
            
            aggregated_results = []
            per_row_keywords = []
            per_row_countries = []
            per_row_languages = []
            seen_urls = set()  # 用于URL去重的集合
            processed_count = 0  # 已处理的结果数量
            for idx, t in enumerate(tasks, 1):
                print(f"\n==== 执行第 {idx} 个任务: {t['keyword']} ====")
                # 临时覆盖搜索配置
                original_num = SEARCH_CONFIG['num_results']
                original_lang = SEARCH_CONFIG['language']
                original_country = SEARCH_CONFIG['country']
                original_exclude = EXCLUDE_TERMS

                try:
                    if t['num_results'] is not None:
                        SEARCH_CONFIG['num_results'] = t['num_results']
                    if t['language'] != '':
                        SEARCH_CONFIG['language'] = t['language']
                    if t['country'] != '':
                        SEARCH_CONFIG['country'] = t['country']
                    if t['exclude']:
                        globals()['EXCLUDE_TERMS'] = t['exclude']

                    tr = process_one_query(t['keyword'], total_target_count=total_target_count, start_index=processed_count)
                    
                    # URL去重：只添加未处理过的URL
                    new_results = []
                    for result in tr:
                        if len(result) >= 4:  # 确保结果包含URL
                            url = result[0]  # URL是第一个元素
                            if url not in seen_urls:
                                seen_urls.add(url)
                                new_results.append(result)
                            else:
                                print(f"  跳过重复URL: {url}")
                    
                    # 只添加去重后的结果
                    aggregated_results.extend(new_results)
                    per_row_keywords.extend([t['keyword']] * len(new_results))
                    per_row_countries.extend([SEARCH_CONFIG['country']] * len(new_results))
                    # 收集每行的搜索语言信息
                    current_language = t['language'] if t['language'] else ''
                    per_row_languages.extend([current_language] * len(new_results))
                    
                    print(f"  本次获取 {len(tr)} 个结果，去重后新增 {len(new_results)} 个结果")
                    # 更新已处理的结果数量
                    processed_count += len(new_results)
                finally:
                    # 还原配置
                    SEARCH_CONFIG['num_results'] = original_num
                    SEARCH_CONFIG['language'] = original_lang
                    SEARCH_CONFIG['country'] = original_country
                    globals()['EXCLUDE_TERMS'] = original_exclude

                if interrupted:
                    break

            # 保存汇总结果
            if aggregated_results:
                from datetime import datetime
                timestamp = datetime.now().strftime('%Y%m%d_%H%M')
                actual_count = len(aggregated_results)
                base = f'mutiple_keywords_{actual_count}个搜索结果'
                docx_filename = f"{base}_{timestamp}.docx"
                html_filename = f"{base}_{timestamp}.html"
                xlsx_filename = f"{base}_{timestamp}.xlsx"

                # 提取唯一的关键词、语言和国家列表
                unique_keywords = list(dict.fromkeys(per_row_keywords))  # 保持顺序的去重
                unique_languages = list(dict.fromkeys([lang for lang in per_row_languages if lang]))  # 去重并过滤空值
                unique_countries = list(dict.fromkeys([country for country in per_row_countries if country]))  # 去重并过滤空值
                
                # 计算总的请求数量（所有任务的总和）
                total_requested_count = sum(t['num_results'] for t in tasks if t['num_results'] is not None)
                
                try:
                    save_results_to_docx(aggregated_results, docx_filename, unique_keywords, unique_languages, unique_countries, total_requested_count)
                    print(f"✓ Word文档已保存到 {docx_filename} 文件中。")
                except Exception as e:
                    print(f"✗ Word文档保存失败: {e}")
                try:
                    save_results_to_html(aggregated_results, html_filename, unique_keywords, unique_languages, unique_countries, total_requested_count)
                    print(f"✓ HTML文件已保存到 {html_filename} 文件中。")
                except Exception as e:
                    print(f"✗ HTML文件保存失败: {e}")
                try:
                    save_results_to_excel(aggregated_results, xlsx_filename, search_keywords_per_row=per_row_keywords, search_countries_per_row=per_row_countries, search_languages_per_row=per_row_languages)
                    print(f"✓ Excel文件已保存到 {xlsx_filename} 文件中。")
                except Exception as e:
                    print(f"✗ Excel文件保存失败: {e}")
            return

    query = SEARCH_KEYWORD
    print(f"搜索关键词: {query}")

    # 显示当前使用的搜索方法
    if SEARCH_METHOD == 'google_api':
        print("使用 Google官方API 进行搜索")
    elif SEARCH_METHOD == 'serp_api':
        print("使用 Serp API 进行搜索")
    else:
        print("使用 googlesearch-python 库进行搜索")
    
    # 在单关键词模式下，total_target_count 就是用户配置的目标数量
    total_target_count = SEARCH_CONFIG['num_results']
    search_results = get_search_results(query, num_results=total_target_count, total_target_count=total_target_count)
    
    # 更新全局状态
    global current_search_results
    current_search_results = search_results
    
    if not search_results:
        print("未能获取任何搜索结果。")
        return

    # 搜索结果已经包含了提取的内容，直接使用
    results = search_results

    # 第二阶段：统一翻译所有结果
    print(f"\n{'='*50}")
    print("第二阶段：批量翻译搜索结果")
    print(f"{'='*50}")
    
    # 检查是否被中断
    if interrupted:
        print("\n⚠ 检测到中断信号，跳过翻译阶段...")
        translated_results = results
    else:
        # 批量翻译所有结果
        translated_results = batch_translate_results(results)
    
    # 更新全局状态
    global current_translated_results
    current_translated_results = translated_results

    # 第三阶段：保存结果
    print(f"\n{'='*50}")
    print("第三阶段：保存结果")
    print(f"{'='*50}")
    
    # 检查是否被中断
    if interrupted:
        print("⚠ 检测到中断信号，跳过正常保存流程...")
        return
    
    # 根据搜索关键词生成文件名
    safe_filename = "".join(c for c in SEARCH_KEYWORD if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_filename = safe_filename.replace(' ', '_')
    # 添加搜索结果数量信息
    actual_count = len(translated_results)
    safe_filename = f"{safe_filename}_{actual_count}个搜索结果"
    # 为最终输出文件名添加时间戳（精确到分钟）
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    docx_filename = f"{safe_filename}_{timestamp}.docx"
    html_filename = f"{safe_filename}_{timestamp}.html"
    xlsx_filename = f"{safe_filename}_{timestamp}.xlsx"
    
    # 保存为Word文档
    try:
        save_results_to_docx(translated_results, docx_filename, requested_count=SEARCH_CONFIG['num_results'])
        print(f"✓ Word文档已保存到 {docx_filename} 文件中。")
    except Exception as e:
        print(f"✗ Word文档保存失败: {e}")
    
    # 保存为HTML文件
    try:
        save_results_to_html(translated_results, html_filename, requested_count=SEARCH_CONFIG['num_results'])
        print(f"✓ HTML文件已保存到 {html_filename} 文件中。")
    except Exception as e:
        print(f"✗ HTML文件保存失败: {e}")
    
    # 保存为Excel文件
    try:
        save_results_to_excel(translated_results, xlsx_filename, SEARCH_KEYWORD, search_country=SEARCH_CONFIG.get('country',''))
        print(f"✓ Excel文件已保存到 {xlsx_filename} 文件中。")
    except Exception as e:
        print(f"✗ Excel文件保存失败: {e}")
    
    # 发送邮件（如果启用）
    if EMAIL_CONFIG['enable_email']:
        try:
            send_email_with_attachment(html_filename, docx_filename, xlsx_filename, num_results=len(translated_results))
            print("✓ 邮件发送成功。")
        except Exception as e:
            print(f"✗ 邮件发送失败: {e}")
    
    return True  # 返回True表示需要最后的暂停

if __name__ == "__main__":
    should_pause = main()
    
    # 标记程序正常完成，避免触发紧急保存
    interrupted = True
    
    # 根据主函数返回值决定是否暂停
    if should_pause:
        # 程序运行结束后暂停，等待用户按键
        print("\n" + "="*60)
        print("程序运行完成！")
        print("="*60)
        input("按回车键退出...")